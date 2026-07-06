# One-time generator: slice the original GDocs DOCX exports into per-template
# DOCX files and replace dotted placeholders / example values with
# {docxtemplater} tags. Source DOCX layout is preserved verbatim.
#
# Usage: python3 scripts/make-templates.py <surat_tugas.docx> <ba_templates.docx>
# Output: templates/docx/{SURAT_TUGAS,BAI,BAKL,BA_PENGUJIAN,BAP}.docx
import re
import sys
import os
from docx import Document

DOTS = re.compile(r'(?:…|\.{4,})[….]*')

OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates', 'docx')


def para_text(p_el):
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    return ''.join(t.text or '' for t in p_el.iter(ns + 't'))


def iter_paragraphs(doc):
    """All paragraphs in document order, including inside tables."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            yield item
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def replace_spans(para, spans):
    """Replace [(start, end, text)] char spans of the paragraph's flat text.

    Flat text = concatenation of run texts. Replacement lands in the first
    overlapping run; the remainder of the span is cleared from later runs, so
    run formatting outside the span is untouched.
    """
    runs = para.runs
    if not runs or not spans:
        return
    for start, end, repl in sorted(spans, key=lambda s: -s[0]):
        pos = 0
        placed = False
        for run in runs:
            rlen = len(run.text)
            rs, re_ = pos, pos + rlen
            pos = re_
            if re_ <= start or rs >= end:
                continue
            ls = max(start, rs) - rs
            le = min(end, re_) - rs
            new = run.text[:ls] + ('' if placed else repl) + run.text[le:]
            placed = True
            run.text = new


def tag_dotted(doc, ordered_tags):
    """Replace every dotted placeholder group, in document order, with tags.

    Adjacent groups separated only by spaces merge into one field.
    """
    idx = 0
    for para in iter_paragraphs(doc):
        flat = ''.join(r.text for r in para.runs)
        matches = list(DOTS.finditer(flat))
        if not matches:
            continue
        merged = []
        for m in matches:
            if merged and flat[merged[-1][1]:m.start()].strip() == '' \
               and m.start() - merged[-1][1] <= 2:
                merged[-1] = (merged[-1][0], m.end())
            else:
                merged.append((m.start(), m.end()))
        spans = []
        for s, e in merged:
            if idx >= len(ordered_tags):
                raise SystemExit(f'more dotted groups than tags near: {flat[:80]!r}')
            spans.append((s, e, '{%s}' % ordered_tags[idx]))
            idx += 1
        replace_spans(para, spans)
    if idx != len(ordered_tags):
        raise SystemExit(f'expected {len(ordered_tags)} dotted groups, found {idx}')


def tag_literals(doc, rules):
    """rules: list of (para_must_contain_or_None, literal_regex, tag)."""
    used = set()
    for para in iter_paragraphs(doc):
        flat = ''.join(r.text for r in para.runs)
        spans = []
        for i, (ctx, pat, tag) in enumerate(rules):
            if ctx is not None and ctx not in flat:
                continue
            for m in re.finditer(pat, flat):
                spans.append((m.start(), m.end(), '{%s}' % tag))
                used.add(i)
        # de-overlap: keep first-listed rule on conflicts
        spans.sort(key=lambda s: (s[0], s[1]))
        pruned, last_end = [], -1
        for s in spans:
            if s[0] >= last_end:
                pruned.append(s)
                last_end = s[1]
        replace_spans(para, pruned)
    missing = [rules[i] for i in range(len(rules)) if i not in used]
    if missing:
        raise SystemExit(f'literal rules not matched: {missing}')


def tag_after_labels(doc, label_tags):
    """Insert tags right after 'Label …… :' occurrences, in document order.

    Labels may share one paragraph, separated by w:br (invisible in flat
    text), so we match sequentially inside each paragraph.
    """
    idx = 0
    for para in iter_paragraphs(doc):
        if idx >= len(label_tags):
            break
        flat = ''.join(r.text for r in para.runs)
        spans = []
        cursor = 0
        while idx < len(label_tags):
            label, tag = label_tags[idx]
            m = re.compile(label + r'[\s\xa0]*:').search(flat, cursor)
            if not m:
                break
            # zero-width insert: rewrite the ':' char as ': {tag}'
            spans.append((m.end() - 1, m.end(), ': {%s}' % tag))
            cursor = m.end()
            idx += 1
        replace_spans(para, spans)
    if idx != len(label_tags):
        raise SystemExit(f'label rules matched {idx}/{len(label_tags)}')


def slice_section(src_path, start_title, end_title):
    """Fresh Document keeping only body elements strictly between the two
    title paragraphs (title paragraphs themselves removed)."""
    doc = Document(src_path)
    body = doc.element.body
    children = list(body)
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    start_i = end_i = None
    for i, el in enumerate(children):
        if el.tag != W + 'p':
            continue
        t = para_text(el).strip()
        if start_i is None and t == start_title:
            start_i = i
        elif start_i is not None and end_i is None and t == end_title:
            end_i = i
            break
    if start_i is None or end_i is None:
        raise SystemExit(f'section not found: {start_title!r} .. {end_title!r} ({start_i},{end_i})')
    for i, el in enumerate(children):
        if el.tag == W + 'sectPr':
            continue
        if i <= start_i or i >= end_i:
            body.remove(el)
    return doc


def main():
    st_path, ba_path = sys.argv[1], sys.argv[2]
    os.makedirs(OUT_DIR, exist_ok=True)

    # ── SURAT TUGAS: standalone file, replace example values ──
    st = Document(st_path)
    tag_literals(st, [
        (None, re.escape('052101/STG/008/SUJBBICON+/2026'), 'nomor'),
        (None, 'Soni\\s+Muhammad Rai', 'nama1'),
        (None, "Bachtiar Rifa.i", 'nama2'),
        (None, 'Hazqy Arfansyah', 'nama3'),
        (None, 'Teknisi', 'jabatanPenerima'),
        (None, 'Distribusi Radio POC', 'uraianTugas'),
        ('Tanggal Tugas', '21\\s+Mei\\s+2026', 'tanggalTugas'),
        ('Bandung,', '21\\s+Mei\\s+2026', 'tanggalSurat'),
        (None, 'GI INDOLIBETY', 'lokasi'),
    ])
    st.save(os.path.join(OUT_DIR, 'SURAT_TUGAS.docx'))
    print('SURAT_TUGAS.docx')

    # ── BAI ──
    bai = slice_section(ba_path, 'TEMPLATE BAI', 'UID JABAR A121601002171')
    tag_dotted(bai, [
        'hari', 'tanggal', 'bulan', 'namaLayanan', 'namaPelanggan',
        'namaLayanan', 'serviceId', 'interface', 'bandwidth', 'originating',
        'terminating', 'noPA',
        'namaPelanggan', 'terminating',
        'namaPerangkat', 'snPerangkat', 'alamatPOP', 'koordinatPOP',
        'namaPerangkatPOP', 'snPOP', 'kanalPort', 'jarakOTDR',
        'namaWakil', 'jabatanWakil', 'alamatKantor', 'kontakWakil',
        'instansiPelanggan', 'ttdPelanggan',
    ])
    bai.save(os.path.join(OUT_DIR, 'BAI.docx'))
    print('BAI.docx')

    # ── BAKL ──
    bakl = slice_section(ba_path, 'TEMPLATE BAKL', 'BAKL A311601001953 MSR')
    tag_dotted(bakl, [
        'hari', 'tanggal', 'bulan', 'namaLayanan', 'noPA',
        'wakilPihakPertama', 'jabatanPihakPertama',
        'instansiPihakKedua', 'wakilPihakKedua', 'jabatanPihakKedua',
        'alamatPihakKedua', 'telpPihakKedua',
        'kendala1', 'kendala2', 'kendala3',
        'lamaTertunda', 'tglMulai', 'tglSelesai',
        'kota', 'tanggalBA',
        'instansiPihakKedua', 'wakilPihakKedua', 'wakilPihakPertama',
    ])
    # the source has a literal year after the date dots — fold it into the tag
    tag_literals(bakl, [(None, r'\{tanggalBA\}\s*2026', 'tanggalBA')])
    bakl.save(os.path.join(OUT_DIR, 'BAKL.docx'))
    print('BAKL.docx')

    # ── BA PENGUJIAN: filled example, replace values ──
    bp = slice_section(ba_path, 'Berita Acara Hasil Pengujian', 'TEMPLATE NODIN')
    tag_literals(bp, [
        (None, re.escape('Dua Ribu Dua Puluh Enam (25-03-2026)'), 'tahun'),
        (None, 'Rabu', 'hari'),
        (None, 'Maret', 'bulan'),
        ('hari ini', r'(?<=\s)25(?=\s)', 'tanggal'),
        (None, 'Saleh Abd Rahman AM', 'namaPihakPertama'),
        (None, 'Direktur', 'jabatanPihakPertama'),
        (None, 'PT Gatra Hita Wasana', 'instansiPihakPertama'),
        (None, r'Jalan Bima Blok TZ No 15.*?15157\.', 'alamatPihakPertama'),
        (None, 'Fajar Sidik Nursyamsi', 'namaPihakKedua'),
        (None, 'TL Delivery Layanan dan B2B dan Ritel Jabar', 'jabatanPihakKedua'),
    ])
    bp.save(os.path.join(OUT_DIR, 'BA_PENGUJIAN.docx'))
    print('BA_PENGUJIAN.docx')

    # ── BAP ──
    bap = slice_section(ba_path, 'TEMPLATE BAP', 'BAP A121201000003')
    tag_dotted(bap, [
        'bulan', 'tahun', 'namaLayanan',
        'tanggalBA', 'instansiPelanggan', 'ttdPelanggan',
    ])
    for para in iter_paragraphs(bap):
        flat = ''.join(r.text for r in para.runs)
        m = re.search(r'(?<=Tanggal)\s*\.\.\s*(?=Bulan)', flat)
        if m:
            replace_spans(para, [(m.start(), m.end(), ' {tanggal} ')])
            break
    else:
        raise SystemExit('BAP tanggal ".." not found')
    tag_after_labels(bap, [
        ('No Sales Order', 'noSalesOrder'),
        ('Nama Pelanggan', 'namaPelanggan'),
        ('Alamat/Lokasi Ori', 'alamatOri'),
        (r'Nama Perangkat\s+&\s+S/N', 'perangkatOri'),
        ('Kanal/Port', 'kanalOri'),
        ('Alamat/Lokasi Ter', 'alamatTer'),
        (r'Nama Perangkat\s+&\s+S/N', 'perangkatTer'),
        ('Kanal/Port', 'kanalTer'),
        ('Kegunaan', 'kegunaan'),
        ('Status Integrasi Originating', 'statusOri'),
        ('Catatan', 'catatanOri'),
        ('Status Integrasi Terminating', 'statusTer'),
        ('Catatan', 'catatanTer'),
        ('Jarak OTDR', 'jarakOTDR'),
    ])
    bap.save(os.path.join(OUT_DIR, 'BAP.docx'))
    print('BAP.docx')


if __name__ == '__main__':
    main()

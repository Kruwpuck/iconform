#!/usr/bin/env python3
"""One-time: insert {tags} into the BAST and NODIN masters (in place).

The masters arrived as plain DOCX with dotted placeholders, not sliced from
the big GDocs export, so make-templates.py doesn't cover them.
"""
import re
import docx

DOTS = re.compile(r'…[….]*|\.{3,}')


def para_sub(p, pairs):
    """Regex/literal replace across the joined runs of a paragraph."""
    if not p.runs:
        return
    text = ''.join(r.text for r in p.runs)
    new = text
    for pat, repl in pairs:
        if isinstance(pat, str):
            new = new.replace(pat, repl)
        else:
            new = pat.sub(repl, new)
    if new != text:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ''


def tag_dotted_cells(table, tags):
    """Replace dotted-only paragraphs in cell order with tags ('' = blank)."""
    # merged cells revisit the same paragraph; once tagged it no longer
    # matches DOTS, so no explicit dedup needed
    it = iter(tags)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if DOTS.fullmatch(p.text.strip()):
                    para_sub(p, [(DOTS, next(it))])


def bast():
    d = docx.Document('templates/docx/BAST.docx')

    for sec in d.sections:
        for p in sec.header.paragraphs:
            para_sub(p, [(re.compile(r'Nomor\s*:.*'), 'Nomor : {nomor}')])

    intro = d.paragraphs[1]
    para_sub(intro, [
        ('Dua Ribu Dua Puluh Lima (...-12-2025)', '{tahun}'),
        ('Desember', '{bulan}'),
    ])
    # remaining dotted runs, in order: hari, tanggal
    order = iter(['{hari}', '{tanggal}'])
    para_sub(intro, [(DOTS, lambda m: next(order))])

    para_sub(d.paragraphs[10], [(DOTS, '{biaya}')])

    tag_dotted_cells(d.tables[0], [
        '{namaPihakPertama}', '{jabatanPihakPertama}', '{instansiPihakPertama}',
        '{berkedudukanPihakPertama}', '',
        '{namaPihakKedua}', '{jabatanPihakKedua}', '{instansiPihakKedua}',
        '{berkedudukanPihakKedua}',
    ])

    d.tables[1].rows[0].cells[2].paragraphs[0].add_run('{perangkat}')
    d.tables[1].rows[1].cells[2].paragraphs[0].add_run('{jumlah}')

    # signature table: drop dotted signer-name lines, keep PIHAK headings
    for row in d.tables[2].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if DOTS.fullmatch(p.text.strip()):
                    para_sub(p, [(DOTS, '')])

    d.save('templates/docx/BAST.docx')


def nodin():
    d = docx.Document('templates/docx/NODIN.docx')

    d.tables[0].rows[0].cells[2].paragraphs[0].add_run('{perihal}')

    order = iter(['{pekerjaan}', '{tim}'])
    para_sub(d.paragraphs[4], [(DOTS, lambda m: next(order))])

    for row in d.tables[1].rows:
        for p in row.cells[0].paragraphs:
            para_sub(p, [(re.compile(r'PRK:\s*, COA :'), 'PRK: {prk} , COA : {coa}')])

    items = d.tables[2]
    # ponytail: single material row; docxtemplater loops need array form data
    # the flat form UI doesn't have — add {#items} loop if multi-row needed
    data = items.rows[2].cells
    for cell, tag in zip(data, ['1', '{material}', '{vol}', '{satuan}', '{hargaSatuan}', '']):
        if tag:
            cell.paragraphs[0].add_run(tag)
    para_sub(items.rows[2].cells[5].paragraphs[0], [('Rp0', 'Rp{jumlahTotal}')])
    para_sub(items.rows[3].cells[5].paragraphs[0], [('Rp0', 'Rp{totalTagihan}')])

    d.save('templates/docx/NODIN.docx')


bast()
nodin()
print('tagged BAST.docx + NODIN.docx')

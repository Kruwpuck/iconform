#!/usr/bin/env python3
"""
scripts/retag.py — one-time script
Copy 8 original DOCX files → templates/docx/<ID>.docx and insert {tag} markers.
Run from repo root: python3 scripts/retag.py
"""

import re
import shutil
from pathlib import Path
from docx import Document

DOTS = re.compile(r'[……\.]{2,}')

SRC = Path('templates')
BA  = SRC / 'Berita Acara'
DST = Path('templates/docx')


# ─── helpers ─────────────────────────────────────────────────────────────────

def merge_runs(para, new_text):
    """Set first run to new_text, clear rest. No-op if already tagged."""
    if '{' in para.text:
        return
    if not para.runs:
        para.add_run(new_text)
    else:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''


def set_run_idx(para, idx, new_text):
    """Replace specific run by index. Caller must guard against re-tagging."""
    if idx < len(para.runs):
        para.runs[idx].text = new_text


def after_colon(para, tag):
    """Replace everything after last ':' with ' {tag}'. No-op if already tagged."""
    if '{' in para.text:
        return
    text = para.text
    pos = text.rfind(':')
    if pos < 0:
        return
    merge_runs(para, text[:pos + 1] + ' ' + tag)


def sub_dots(para, *tags):
    """Replace DOTS sequences with tags in order. No-op if already tagged."""
    if '{' in para.text:
        return
    text = para.text
    for tag in tags:
        text = DOTS.sub(tag, text, count=1)
    merge_runs(para, text)


def cell_after_colon(cell, para_idx, tag):
    after_colon(cell.paragraphs[para_idx], tag)


def cell_sub_dots(cell, para_idx, *tags):
    sub_dots(cell.paragraphs[para_idx], *tags)


def cell_merge(cell, para_idx, new_text):
    merge_runs(cell.paragraphs[para_idx], new_text)


def add_img_para(para, tag):
    """Write image tag into an empty paragraph."""
    if para.text.strip() or ('{' in para.text):
        return
    if para.runs:
        para.runs[0].text = tag
    else:
        para.add_run(tag)


def add_img_cell(cell, para_idx, tag):
    """Add image tag to cell paragraph; adds new paragraph if index doesn't exist."""
    while len(cell.paragraphs) <= para_idx:
        cell.add_paragraph()
    add_img_para(cell.paragraphs[para_idx], tag)


# ─── SURAT_TUGAS ─────────────────────────────────────────────────────────────

def do_surat_tugas():
    dst = DST / 'SURAT_TUGAS.docx'
    shutil.copy(SRC / 'TEMPLATE_Surat Tugas_Nama.docx', dst)
    doc = Document(dst)
    p = doc.paragraphs

    merge_runs(p[10], 'Nomor : {nomor}')
    merge_runs(p[13], 'Nama\t: {namaPemberi}')
    merge_runs(p[14], 'Jabatan\t: {jabatanPemberi}')
    merge_runs(p[16], 'Nama\t: -\tSdr. {nama1}')
    merge_runs(p[17], ' Sdr.\t{nama2}')
    merge_runs(p[18], ' Sdr.\t\t\t{nama3}')
    merge_runs(p[19], 'Jabatan\t: {jabatanPenerima}')
    merge_runs(p[22], 'Uraian Tugas\t: {uraianTugas}')
    # p23 runs: ['Tanggal Tugas\t  : ', ' ', 's/d ']
    if '{' not in p[23].text:
        set_run_idx(p[23], 1, '{tanggalTugas} ')
        set_run_idx(p[23], 2, 's/d {tanggalTugasSelesai}')
    # p24 runs: ['Lokasi', '\t\t  ', ': ']
    if '{' not in p[24].text:
        set_run_idx(p[24], 2, ': {lokasi}')
    # p28 runs: ['Bandung, ', '21 Mei 2026', '\n', 'ENGINEER...', ...]
    if '{' not in p[28].text:
        set_run_idx(p[28], 1, '{tanggalSurat}')
    merge_runs(p[34], '{namaPemberi}')
    # TTD/stempel in blank paragraphs between header (p28) and name (p34)
    add_img_para(p[29], '{%ttd}')
    add_img_para(p[30], '{%stempel}')

    doc.save(dst)
    print(f'  ✓ SURAT_TUGAS')


# ─── BAI ─────────────────────────────────────────────────────────────────────

def _tag_bai_body(doc):
    """Tag BAI/UID_JABAR body paragraphs with dots → {tags}."""
    p = doc.paragraphs
    # Simple single-tag after colon
    for idx, tag in [
        (3,  '{namaLayanan}'),
        (4,  '{serviceId}'),
        (5,  '{interface}'),
        (6,  '{bandwidth}'),
        (7,  '{originating}'),
        (8,  '{terminating}'),
        (9,  '{noPA}'),
        (11, '{namaPelanggan}'),
        # p12 = extra Terminating row, no field — skip
        (15, '{alamatPOP}'),
        (16, '{koordinatPOP}'),
        (18, '{kanalPort}'),
        (19, '{jarakOTDR}'),
        (22, '{namaWakil}'),
        (23, '{jabatanWakil}'),
        (24, '{alamatKantor}'),
        (25, '{kontakWakil}'),
    ]:
        after_colon(p[idx], tag)

    # p14/p17: two tags (namaPerangkat + SN) per paragraph
    for idx, t1, t2 in [
        (14, '{namaPerangkat}', '{snPerangkat}'),
        (17, '{namaPerangkatPOP}', '{snPOP}'),
    ]:
        sub_dots(p[idx], t1, t2)

    merge_runs(p[28], 'Nama Project Team Leader : {namaProjectLeader}')


def _tag_bai_sig(doc):
    """Tag BAI/UID_JABAR signature table."""
    tbl = doc.tables[0]
    cell_sub_dots(tbl.rows[0].cells[0], 0, '{instansiPelanggan}')
    cell_sub_dots(tbl.rows[2].cells[0], 0, '{namaWakil}')
    merge_runs(tbl.rows[2].cells[1].paragraphs[0], '  ( {namaProjectLeader} )')
    # TTD/stempel in PLN cell row[1]
    pln_cell = tbl.rows[1].cells[1]
    add_img_cell(pln_cell, 0, '{%ttd}')
    add_img_cell(pln_cell, 1, '{%stempel}')


def do_bai():
    dst = DST / 'BAI.docx'
    shutil.copy(BA / 'TEMPLATE_Berita Acara Instalasi_Surat BAI_NO. PA.docx', dst)
    doc = Document(dst)

    # p1: 'Pada hari ini ………. Tanggal …. Bulan …………… Tahun Dua Ribu Dua Puluh Enam ...'
    p1 = doc.paragraphs[1]
    if '{' not in p1.text:
        text = p1.text
        text = re.sub(r'(hari ini )[…\.]+', r'\1{hari}', text, count=1)
        text = re.sub(r'(Tanggal )[…\.]+', r'\1{tanggal}', text, count=1)
        text = re.sub(r'(Bulan )[…\.]+', r'\1{bulan}', text, count=1)
        text = text.replace('Dua Ribu Dua Puluh Enam', '{tahun}', 1)
        merge_runs(p1, text)

    _tag_bai_body(doc)
    _tag_bai_sig(doc)
    doc.save(dst)
    print('  ✓ BAI')


# ─── UID_JABAR (BAPB) ────────────────────────────────────────────────────────

def do_uid_jabar():
    dst = DST / 'UID_JABAR.docx'
    shutil.copy(BA / 'TEMPLATE_Surat Berita Acara Pengujian Barang_BAPB_NO. PA.docx', dst)
    doc = Document(dst)

    # p1: 'Pada hari ini Rabu Tanggal 25 Bulan Maret Tahun Dua Ribu Dua Puluh Enam ...'
    p1 = doc.paragraphs[1]
    if '{' not in p1.text:
        text = p1.text
        text = re.sub(r'(hari ini )\w+', r'\1{hari}', text, count=1)
        text = re.sub(r'(Tanggal )\d+', r'\1{tanggal}', text, count=1)
        text = re.sub(r'(Bulan )\w+', r'\1{bulan}', text, count=1)
        text = re.sub(r'Tahun Dua Ribu Dua Puluh \w+', 'Tahun {tahun}', text, count=1)
        merge_runs(p1, text)

    # Body: some fields are dotted (same as BAI), some are literals
    paras = doc.paragraphs
    after_colon(paras[3],  '{namaLayanan}')   # I-WIN literal
    after_colon(paras[4],  '{serviceId}')      # 121610001072
    after_colon(paras[5],  '{interface}')      # dots
    after_colon(paras[6],  '{bandwidth}')      # dots
    after_colon(paras[7],  '{originating}')    # JL. ASIA AFRIKA literal
    after_colon(paras[8],  '{terminating}')    # GEDUNG BALAI literal
    after_colon(paras[9],  '{noPA}')           # A121610001700
    after_colon(paras[11], '{namaPelanggan}')  # PT. PLN (PERSERO) ... literal
    # p12 = extra Terminating row — skip
    # p14: 'Nama Perangkat\t: Existing\tSN: Existing'
    if '{' not in paras[14].text:
        t = paras[14].text.replace('Existing', '{namaPerangkat}', 1).replace('Existing', '{snPerangkat}', 1)
        merge_runs(paras[14], t)
    after_colon(paras[15], '{alamatPOP}')      # Existing
    after_colon(paras[16], '{koordinatPOP}')   # Existing
    # p17: 'Nama Perangkat POP\t: Existing\tSN: Existing'
    if '{' not in paras[17].text:
        t = paras[17].text.replace('Existing', '{namaPerangkatPOP}', 1).replace('Existing', '{snPOP}', 1)
        merge_runs(paras[17], t)
    after_colon(paras[18], '{kanalPort}')      # Existing
    after_colon(paras[19], '{jarakOTDR}')      # Existing
    after_colon(paras[22], '{namaWakil}')
    after_colon(paras[23], '{jabatanWakil}')
    after_colon(paras[24], '{alamatKantor}')
    after_colon(paras[25], '{kontakWakil}')
    merge_runs(paras[28], 'Nama Project Team Leader : {namaProjectLeader}')

    # Sig table
    tbl = doc.tables[0]
    # [0,0] = 'PT. PLN (PERSERO) UNIT INDUK DISTRIBUSI JAWA BARAT' → {instansiPelanggan}
    merge_runs(tbl.rows[0].cells[0].paragraphs[0], '{instansiPelanggan}')
    cell_sub_dots(tbl.rows[2].cells[0], 0, '{namaWakil}')
    merge_runs(tbl.rows[2].cells[1].paragraphs[0], '  ( {namaProjectLeader} )')
    pln_cell = tbl.rows[1].cells[1]
    add_img_cell(pln_cell, 0, '{%ttd}')
    add_img_cell(pln_cell, 1, '{%stempel}')

    doc.save(dst)
    print('  ✓ UID_JABAR')


# ─── BAKL ────────────────────────────────────────────────────────────────────

def do_bakl():
    dst = DST / 'BAKL.docx'
    shutil.copy(BA / 'TEMPLATE_BERITA ACARA KENDALA LAPANGAN_BAKL_NO PA.docx', dst)
    doc = Document(dst)
    p = doc.paragraphs

    # p3: 'Pada hari ini …… Tanggal …… Bulan …… Tahun 2026telah ... layanan …… '
    if '{' not in p[3].text:
        text = p[3].text
        text = DOTS.sub('{hari}',       text, count=1)
        text = DOTS.sub('{tanggal}',    text, count=1)
        text = DOTS.sub('{bulan}',      text, count=1)
        text = text.replace('2026telah', '{tahun} telah', 1)
        text = DOTS.sub('{namaLayanan}', text, count=1)
        merge_runs(p[3], text)

    # p5: 'No PA: ………………………….    '
    after_colon(p[5], '{noPA}')

    # p9-11: pure dots → kendala1-3
    merge_runs(p[9],  '{kendala1}')
    merge_runs(p[10], '{kendala2}')
    merge_runs(p[11], '{kendala3}')

    # p13: 'Yang diperkirakan akan tertunda selama ….. hari, sejak tanggal …….. s/d ……. ...'
    sub_dots(p[13], '{lamaTertunda}', '{tglMulai}', '{tglSelesai}')

    # p19: dots-comma-dots-tahun → '{kota}, {tanggalBA}'
    merge_runs(p[19], '{kota}, {tanggalBA}')

    # table[0]: pihak rows
    t0 = doc.tables[0]
    # PLN side (col 2): diwakili [1,2], jabatan [2,2] — PLN address/phone kept
    cell_sub_dots(t0.rows[1].cells[2], 0, '{wakilPihakPertama}')
    cell_sub_dots(t0.rows[2].cells[2], 0, '{jabatanPihakPertama}')
    # Pihak Kedua (rows 6-10, col 2)
    cell_sub_dots(t0.rows[6].cells[2],  0, '{instansiPihakKedua}')
    cell_sub_dots(t0.rows[7].cells[2],  0, '{wakilPihakKedua}')
    cell_sub_dots(t0.rows[8].cells[2],  0, '{jabatanPihakKedua}')
    cell_sub_dots(t0.rows[9].cells[2],  0, '{alamatPihakKedua}')
    cell_sub_dots(t0.rows[10].cells[2], 0, '{telpPihakKedua}')

    # table[2]: sig
    t2 = doc.tables[2]
    cell_sub_dots(t2.rows[0].cells[0], 0, '{instansiPihakKedua}')
    cell_sub_dots(t2.rows[2].cells[0], 0, '{wakilPihakKedua}')
    cell_sub_dots(t2.rows[2].cells[2], 0, '{wakilPihakPertama}')
    # TTD/stempel in PLN cell [1,2]
    pln_cell = t2.rows[1].cells[2]
    add_img_cell(pln_cell, 0, '{%ttd}')
    add_img_cell(pln_cell, 1, '{%stempel}')

    doc.save(dst)
    print('  ✓ BAKL')


# ─── BA_PENGUJIAN (BAHP) ─────────────────────────────────────────────────────

def do_ba_pengujian():
    dst = DST / 'BA_PENGUJIAN.docx'
    shutil.copy(BA / 'TEMPLATE_BERITA ACARA HASIL PENGUJIAN_BAHP_Nama Instansi Pihak I.docx', dst)
    doc = Document(dst)

    # Header: p1 run[1] = literal nomor
    hdr = doc.sections[0].header
    h1 = hdr.paragraphs[1]
    if '{' not in h1.text and len(h1.runs) >= 2:
        h1.runs[1].text = '{nomor}'
    # Add {%logoMitra} as new header paragraph
    logo_para = hdr.add_paragraph()
    logo_para.add_run('{%logoMitra}')

    # p1 runs: ['Pada hari ini,', ' Rabu ', 'tanggal ', '25 ', 'bulan ', 'Maret ', 'tahun ', 'Dua Ribu... (25-03-2026)', ', kami...']
    p1 = doc.paragraphs[1]
    if '{' not in p1.text:
        set_run_idx(p1, 1, ' {hari} ')
        set_run_idx(p1, 3, '{tanggal} ')
        set_run_idx(p1, 5, '{bulan} ')
        set_run_idx(p1, 7, '{tahun}')

    # table[0]: pihak identification
    t0 = doc.tables[0]
    # Pihak Pertama (rows 0-3): [0,3] p0=nama, p1=jabatan; [1,3]=instansi; [2,3] p0=alamat
    r0c3 = t0.rows[0].cells[3]
    merge_runs(r0c3.paragraphs[0], '{namaPihakPertama}')
    merge_runs(r0c3.paragraphs[1], '{jabatanPihakPertama}')
    cell_merge(t0.rows[1].cells[3], 0, '{instansiPihakPertama}')
    cell_merge(t0.rows[2].cells[3], 0, '{alamatPihakPertama}')
    # Pihak Kedua (rows 5-6): [5,3]=nama; [6,3]=jabatan
    cell_merge(t0.rows[5].cells[3], 0, '{namaPihakKedua}')
    cell_merge(t0.rows[6].cells[3], 0, '{jabatanPihakKedua}')

    # table[1]: sig
    t1 = doc.tables[1]
    cell_merge(t1.rows[2].cells[0], 0, '{namaPihakPertama}')
    cell_merge(t1.rows[2].cells[2], 0, '{namaPihakKedua}')
    # TTD/stempel in PLN (Pihak Kedua) cell [1,2]
    pln_cell = t1.rows[1].cells[2]
    add_img_cell(pln_cell, 0, '{%ttd}')
    add_img_cell(pln_cell, 1, '{%stempel}')

    doc.save(dst)
    print('  ✓ BA_PENGUJIAN')


# ─── BAP ─────────────────────────────────────────────────────────────────────

def do_bap():
    dst = DST / 'BAP.docx'
    shutil.copy(BA / 'TEMPLATE_BERITA ACARA PEMAKAIAN_BAP_No Sales Order.docx', dst)
    doc = Document(dst)
    p = doc.paragraphs

    # p2 runs: [..., ' .. ', 'Bulan ', '……… ', 'Tahun ', '……………………. ', ..., '………………………….  ', ...]
    if '{' not in p[2].text:
        set_run_idx(p[2], 3, ' {tanggal} ')
        set_run_idx(p[2], 5, '{bulan} ')
        set_run_idx(p[2], 7, '{tahun} ')
        set_run_idx(p[2], 9, '{namaLayanan}  ')

    # p6 runs: ['     \t', 'Bandung, ………………………………..']
    if '{' not in p[6].text:
        set_run_idx(p[6], 1, 'Bandung, {tanggalBA}')

    # table[0]: complex single-column merged layout
    t0 = doc.tables[0]
    # [0,0] p0: 'No Sales Order\t: \nNama Pelanggan\t: ' (embedded newline, one run)
    c00 = t0.rows[0].cells[0]
    if '{' not in c00.paragraphs[0].text:
        merge_runs(c00.paragraphs[0],
                   'No Sales Order\t: {noSalesOrder}\nNama Pelanggan\t: {namaPelanggan}')
    cell_after_colon(c00, 1, '{alamatOri}')
    cell_after_colon(c00, 2, '{perangkatOri}')
    cell_after_colon(c00, 3, '{kanalOri}')

    c10 = t0.rows[1].cells[0]
    cell_after_colon(c10, 0, '{alamatTer}')
    cell_after_colon(c10, 1, '{perangkatTer}')
    cell_after_colon(c10, 2, '{kanalTer}')

    cell_after_colon(t0.rows[2].cells[0], 0, '{kegunaan}')

    c30 = t0.rows[3].cells[0]
    cell_after_colon(c30, 0, '{statusOri}')
    cell_after_colon(c30, 1, '{catatanOri}')
    cell_after_colon(c30, 2, '{statusTer}')
    cell_after_colon(c30, 3, '{catatanTer}')

    cell_after_colon(t0.rows[4].cells[0], 0, '{jarakOTDR}')

    # table[1]: sig
    t1 = doc.tables[1]
    pelanggan_cell = t1.rows[0].cells[0]
    pln_cell       = t1.rows[0].cells[1]
    cell_sub_dots(pelanggan_cell, 0, '{instansiPelanggan}')
    cell_sub_dots(pelanggan_cell, 5, '{namaWakilPelanggan}')
    merge_runs(pln_cell.paragraphs[5], '{namaWakilIcon}')
    # TTD/stempel in PLN cell (right side)
    add_img_cell(pln_cell, 1, '{%ttd}')
    add_img_cell(pln_cell, 2, '{%stempel}')

    doc.save(dst)
    print('  ✓ BAP')


# ─── BAST ────────────────────────────────────────────────────────────────────

def do_bast():
    dst = DST / 'BAST.docx'
    shutil.copy(BA / 'TEMPLATE_BERITA ACARA SERAH TERIMA BARANG_BAST_Perangkat.docx', dst)
    doc = Document(dst)

    # Header: p1 run[1] = literal nomor
    hdr = doc.sections[0].header
    h1 = hdr.paragraphs[1]
    if '{' not in h1.text and len(h1.runs) >= 2:
        h1.runs[1].text = '{nomor}'

    # p1 runs: ['Pada hari ini,', ' …….. ', 'tanggal ', '… ', 'bulan ', 'Desember ', 'tahun ', 'Dua Ribu Dua Puluh Lima (...-12-2025)', ', kami...']
    p1 = doc.paragraphs[1]
    if '{' not in p1.text:
        set_run_idx(p1, 1, ' {hari} ')
        set_run_idx(p1, 3, '{tanggal} ')
        set_run_idx(p1, 5, '{bulan} ')
        set_run_idx(p1, 7, '{tahun}')

    # p10: '...biaya …………. oleh PIHAK...'
    sub_dots(doc.paragraphs[10], '{biaya}')

    # table[0]: pihak identification
    t0 = doc.tables[0]
    r0c3 = t0.rows[0].cells[3]
    cell_sub_dots(r0c3, 0, '{namaPihakPertama}')
    cell_sub_dots(r0c3, 1, '{jabatanPihakPertama}')
    cell_sub_dots(t0.rows[1].cells[3], 0, '{instansiPihakPertama}')
    cell_sub_dots(t0.rows[2].cells[3], 0, '{berkedudukanPihakPertama}')
    cell_sub_dots(t0.rows[5].cells[3], 0, '{namaPihakKedua}')
    cell_sub_dots(t0.rows[6].cells[3], 0, '{jabatanPihakKedua}')
    cell_sub_dots(t0.rows[7].cells[3], 0, '{instansiPihakKedua}')
    cell_sub_dots(t0.rows[8].cells[3], 0, '{berkedudukanPihakKedua}')

    # table[1]: perangkat/jumlah
    t1 = doc.tables[1]
    cell_merge(t1.rows[0].cells[1], 0, ': {perangkat}')
    cell_merge(t1.rows[1].cells[1], 0, ': {jumlah}')

    # table[2]: sig
    t2 = doc.tables[2]
    cell_sub_dots(t2.rows[2].cells[0], 0, '{namaPihakPertama}')
    cell_sub_dots(t2.rows[2].cells[2], 0, '{namaPihakKedua}')
    pln_cell = t2.rows[1].cells[2]
    add_img_cell(pln_cell, 0, '{%ttd}')
    add_img_cell(pln_cell, 1, '{%stempel}')

    doc.save(dst)
    print('  ✓ BAST')


# ─── NODIN ───────────────────────────────────────────────────────────────────

def do_nodin():
    dst = DST / 'NODIN.docx'
    shutil.copy(BA / 'TEMPLATE_NODIN_Nama Agenda.docx', dst)
    doc = Document(dst)

    # table[0]: perihal in cell[0,2] (col0=label, col1=':', col2=value)
    t0p = doc.tables[0].rows[0].cells[2].paragraphs[0]
    if '{' not in t0p.text:
        if t0p.runs:
            t0p.runs[0].text = '{perihal}'
        else:
            t0p.add_run('{perihal}')

    # p4: 'Sehubungan telah dilaksanakannya ….., yang dilaksanakan oleh Tim dari ……., Maka ...'
    p4 = doc.paragraphs[4]
    if '{' not in p4.text:
        text = p4.text
        text = re.sub(r'[…\.]+(?=,)', '{pekerjaan}', text, count=1)
        text = re.sub(r'(dari )[…\.]+', r'\1{tim}', text, count=1)
        merge_runs(p4, text)

    # table[1]: MI Dana, MI List, PRK/COA
    t1 = doc.tables[1]
    if '{' not in t1.rows[0].cells[0].paragraphs[0].text:
        t1.rows[0].cells[0].paragraphs[0].runs[0].text = '{miDana}'
    if '{' not in t1.rows[1].cells[0].paragraphs[0].text:
        t1.rows[1].cells[0].paragraphs[0].runs[0].text = '{miList}'
    if '{' not in t1.rows[2].cells[0].paragraphs[0].text:
        t1.rows[2].cells[0].paragraphs[0].runs[0].text = 'PRK: {prk} , COA : {coa}'

    # table[2] row2: loop template row
    t2 = doc.tables[2]
    loop_tags = ['{#items}', '{material}', '{vol}', '{satuan}', '{hargaSatuan}', 'Rp{jumlahTotal}{/items}']
    for ci, tag in enumerate(loop_tags):
        cell = t2.rows[2].cells[ci]
        para = cell.paragraphs[0]
        if '{' not in para.text:
            if para.runs:
                para.runs[0].text = tag
            else:
                para.add_run(tag)

    # table[2] row3: total tagihan
    total_cell = t2.rows[3].cells[5]
    tp = total_cell.paragraphs[0]
    if tp.text.strip() == 'Rp0':
        tp.runs[0].text = 'Rp{totalTagihan}'

    doc.save(dst)
    print('  ✓ NODIN')


# ─── main ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('Re-tagging original templates...')
    do_surat_tugas()
    do_bai()
    do_uid_jabar()
    do_bakl()
    do_ba_pengujian()
    do_bap()
    do_bast()
    do_nodin()
    print('Done. Run tag-dump to verify.')

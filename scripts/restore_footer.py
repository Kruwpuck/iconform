#!/usr/bin/env python3
"""Restore the SURAT_TUGAS bottom decorative bar ("Shape1").

The original master (templates/TEMPLATE_Surat Tugas_Nama.docx) carries TWO
floating anchored drawings in the document body — there is no word/footer*.xml
part anywhere in this repo, the bars are body-anchored art:

  1. "image1.jpg"  positionV relativeFrom=page   offset 0         -> top bar
  2. "Shape1"      positionV relativeFrom=margin offset 9801860   -> bottom bar

Anchor #2 was dropped from templates/docx/SURAT_TUGAS.docx somewhere in the
retag/revisi pipeline (see scripts/retag*.py, scripts/revisi*.py), which is why
generated Surat Tugas documents render without the footer colour bar.

This copies the master's Shape1 run (mc:AlternateContent -> wpg group of three
custGeom vector shapes filled #16a2b9 / #f5ec12 / #155a73, plus a VML
mc:Fallback) into the working template's last content paragraph, mirroring the
master's own placement. The working template is NOT regenerated from the master
— many rounds of tag/layout edits live in it and must survive.

Idempotent — safe to re-run. Run from repo root:  python scripts/restore_footer.py
"""
import copy
import zipfile
from lxml import etree
import docx
from docx.oxml.ns import qn

MASTER = 'templates/TEMPLATE_Surat Tugas_Nama.docx'
TARGET = 'templates/docx/SURAT_TUGAS.docx'
SHAPE = 'Shape1'
WP = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
MC = '{http://schemas.openxmlformats.org/markup-compatibility/2006}'
# mirror of the master: the drawing run sits in the last content paragraph,
# the one holding the signer's name (literal name in the master, tagged here).
ANCHOR_PARA_TAG = '{namaPemberi}'


def shape1_run():
    """Deep-copy the master's Shape1 <w:r> out of a read-only zip handle.

    The master is reference material: opened read-only, never written.
    """
    with zipfile.ZipFile(MASTER) as z:
        root = etree.fromstring(z.read('word/document.xml'))
    docPr = next(d for d in root.iter(WP + 'docPr') if d.get('name') == SHAPE)
    # climb to the enclosing w:r (docPr -> anchor -> drawing -> mc:Choice ->
    # mc:AlternateContent -> w:r)
    run = enclosing_run(docPr)
    assert run.find(MC + 'AlternateContent') is not None, 'expected AlternateContent wrapper'
    return copy.deepcopy(run)


def used_docpr_ids(part_root):
    return {int(d.get('id')) for d in part_root.iter(WP + 'docPr') if d.get('id')}


def enclosing_run(el):
    while el.tag != qn('w:r'):
        el = el.getparent()
    return el


def restore(d):
    body = d.element.body
    # The signature block, not the "Nama: {namaPemberi}" line near the top —
    # {namaPemberi} appears twice, and only the LAST one is the master's
    # equivalent paragraph.
    target_p = [p for p in d.paragraphs if ANCHOR_PARA_TAG in p.text][-1]

    if any(x.get('name') == SHAPE for x in target_p._p.iter(WP + 'docPr')):
        print(f'  {SHAPE} already in the target paragraph - nothing to do')
        return False

    # Converge from any state: drop a copy sitting in the wrong paragraph
    # rather than leaving a duplicate behind.
    for docPr in [x for x in body.iter(WP + 'docPr') if x.get('name') == SHAPE]:
        run = enclosing_run(docPr)
        run.getparent().remove(run)
        print(f'  removed a misplaced {SHAPE} run')

    run = shape1_run()

    # The empty spacer textbox references pStyle "normal1", which exists in the
    # master's styles.xml but NOT in the target's. Word tolerates a dangling
    # pStyle, but the textbox is blank so the reference buys nothing — drop it
    # rather than importing a whole style definition for invisible content.
    for pStyle in list(run.iter(qn('w:pStyle'))):
        if pStyle.get(qn('w:val')) == 'normal1':
            pStyle.getparent().remove(pStyle)

    # wp:docPr/@id must be unique per document; the master numbered Shape1 as
    # id=2 and the target already uses id=2 for its top bar.
    taken = used_docpr_ids(body)
    for docPr in run.iter(WP + 'docPr'):
        if int(docPr.get('id')) in taken:
            new_id = max(taken) + 1
            docPr.set('id', str(new_id))
            print(f'  remapped docPr id -> {new_id} (2 was taken by the top bar)')

    # insert as the paragraph's first run, exactly as in the master
    pPr = target_p._p.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(run)
    else:
        target_p._p.insert(0, run)
    print(f'  inserted {SHAPE} run into paragraph {ANCHOR_PARA_TAG!r}')
    return True


def report(path):
    """Structural verification — there is no local renderer, so read the XML
    back out of the saved file and assert on it."""
    with zipfile.ZipFile(path) as z:
        bad = z.testzip()
        assert bad is None, f'corrupt zip entry: {bad}'
        root = etree.fromstring(z.read('word/document.xml'))

    anchors = root.findall('.//' + WP + 'anchor')
    print(f'  anchors: {len(anchors)}')
    for a in anchors:
        name = a.find(WP + 'docPr').get('name')
        pv = a.find(WP + 'positionV')
        ext = a.find(WP + 'extent')
        print(f'    - {name!r} id={a.find(WP + "docPr").get("id")} '
              f'positionV relativeFrom={pv.get("relativeFrom")} '
              f'offset={pv.find(WP + "posOffset").text} '
              f'extent={ext.get("cx")}x{ext.get("cy")}')

    # every prefix the document uses must be declared
    declared = set(root.nsmap.values())
    for el in root.iter():
        if isinstance(el.tag, str) and el.tag.startswith('{'):
            uri = el.tag.split('}')[0][1:]
            assert uri in declared or uri in el.nsmap.values(), f'undeclared ns {uri}'

    d = docx.Document(path)  # python-docx must still parse it
    texts = [p.text for p in d.paragraphs]
    return anchors, texts


TAGS = ['{#petugas}', '{/petugas}', '{nama}', '{namaPemberi}', '{jabatanPemberi}',
        '{jabatanPenerima}', '{uraianTugas}', '{tanggalTugas}',
        '{tanggalTugasSelesai}', '{lokasi}', '{%ttd}', '{%stempel}',
        '{nomor}', '{tanggalSurat}']


def main():
    print(f'before: {TARGET}')
    before_anchors, before_texts = report(TARGET)
    before_tags = {t: sum(p.count(t) for p in before_texts) for t in TAGS}

    d = docx.Document(TARGET)
    changed = restore(d)
    if changed:
        d.save(TARGET)
        print('  saved')

    print(f'after: {TARGET}')
    after_anchors, after_texts = report(TARGET)
    after_tags = {t: sum(p.count(t) for p in after_texts) for t in TAGS}

    assert len(after_anchors) == 2, f'expected 2 anchors, got {len(after_anchors)}'
    names = [a.find(WP + 'docPr').get('name') for a in after_anchors]
    assert names[1] == SHAPE, f'second anchor should be {SHAPE}, got {names[1]}'
    bar = after_anchors[1]
    assert bar.find(WP + 'positionV').get('relativeFrom') == 'margin'
    assert bar.find(WP + 'positionV/' + WP + 'posOffset').text == '9801860'
    assert bar.find(WP + 'extent').get('cx') == '5340985'
    assert bar.find(WP + 'extent').get('cy') == '914400'
    assert before_tags == after_tags, f'tag counts changed:\n{before_tags}\n{after_tags}'
    assert before_texts == after_texts, 'paragraph text changed'
    print('  tag counts unchanged:', after_tags)
    print('OK')


if __name__ == '__main__':
    main()

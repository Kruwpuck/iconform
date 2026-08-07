"""Fix BAI spacing — tab stops + hanging indent for long-value wrapping.

- Uniform tab stop at 2880 twips so all colons line up in one column.
- Hanging indent (left=2880 hanging=2880) on the long value rows
  (Originating / Terminating): a wrapped address folds back under the
  value column instead of the left margin. Hanging pos == tab stop, so
  the colon stays put.
- Ensures a second "Terminating" row exists right after "Nama Pelanggan"
  (mirrors the customer master layout — image 3). Idempotent.
"""

import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TPL = Path(__file__).parent.parent / "templates" / "docx"

# BAI + UID_JABAR share the same spesifikasi layout
FILES = ["BAI.docx", "UID_JABAR.docx"]

# Rows whose value can be a long address → need hanging indent so wraps align.
WRAP_TAGS = ("{originating}", "{terminating}")


def fix_file(name):
    doc = Document(TPL / name)

    for p in doc.paragraphs:
        text = "".join(r.text or "" for r in p.runs).strip()
        # match any data-field paragraph: has a tab+colon pattern
        if '\t:' not in text and ('\t' not in text or ':' not in text.split('\t', 1)[-1][:3]):
            continue

        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p._element.insert(0, pPr)

        # Reset tabs
        existing_tabs = pPr.find(qn('w:tabs'))
        if existing_tabs is not None:
            pPr.remove(existing_tabs)

        # Reset indent — we re-add a correct hanging indent only where needed.
        existing_ind = pPr.find(qn('w:ind'))
        if existing_ind is not None:
            pPr.remove(existing_ind)

        # Long-value rows get a hanging indent so wrapped lines align under the
        # value column (2880). Hanging position == tab stop, colons stay fixed.
        if any(tag in text for tag in WRAP_TAGS):
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '2880')
            ind.set(qn('w:hanging'), '2880')
            pPr.append(ind)

        # Uniform tab at 2880 twips (2 inch) — fits the longest label.
        # Rows with a second tab ("...: {value}\tSN: {sn}") get a stop at 5760
        # so the SN column lines up regardless of value length.
        n_tabs = sum(1 for r in p.runs for ch in r._r if ch.tag == qn('w:tab'))
        tabs = OxmlElement('w:tabs')
        for pos in (['2880', '5760'] if n_tabs >= 2 else ['2880']):
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'left')
            tab.set(qn('w:pos'), pos)
            tabs.append(tab)
        pPr.append(tabs)

    ensure_second_terminating(doc)

    doc.save(TPL / name)
    print(f"{name}: tab stops 2880, hanging indent on wrap rows, 2nd Terminating ensured")


def ensure_second_terminating(doc):
    """Insert a Terminating row right after Nama Pelanggan if missing."""
    term_para = None
    pelanggan_para = None
    term_count = 0
    for p in doc.paragraphs:
        text = "".join(r.text or "" for r in p.runs)
        if '{terminating}' in text:
            term_count += 1
            if term_para is None:
                term_para = p
        if '{namaPelanggan}' in text:
            pelanggan_para = p

    if term_count >= 2 or term_para is None or pelanggan_para is None:
        return  # already present, or nothing to clone from

    clone = copy.deepcopy(term_para._element)
    pelanggan_para._element.addnext(clone)


if __name__ == '__main__':
    for f in FILES:
        fix_file(f)

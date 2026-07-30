#!/usr/bin/env python3
"""Revisi BAST: stack Perangkat/Jumlah top-bottom, trim excess blank
paragraphs before signature (was overflowing to page 2), nudge header
logo right so it stops colliding with the centered title/nomor lines.
Run once from repo root:  python scripts/revisi_bast.py
"""
import docx
from docx.oxml.ns import qn

PATH = 'templates/docx/BAST.docx'


def fix_perangkat_jumlah(d):
    t = d.tables[1]
    if len(t.rows) > 1:
        return  # already stacked, idempotent
    row0 = t.rows[0]
    if len(row0.cells) < 4:
        return
    jumlah_label = row0.cells[2].text
    jumlah_value = row0.cells[3].text
    # blank out the side-by-side Jumlah cells on row 0
    row0.cells[2].paragraphs[0].runs and None
    for cell in (row0.cells[2], row0.cells[3]):
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                pass
    new_row = t.add_row()
    new_row.cells[0].paragraphs[0].add_run(jumlah_label)
    new_row.cells[1].paragraphs[0].add_run(jumlah_value)


def trim_blank_paragraphs(d):
    # find the run of blank paragraphs right before the signature table
    # (directly following "Demikian Berita Acara...")
    blanks = []
    found_demikian = False
    for p in d.paragraphs:
        if 'Demikian Berita Acara' in p.text:
            found_demikian = True
            continue
        if found_demikian:
            if p.text.strip() == '':
                blanks.append(p)
            else:
                break
    keep = 2
    for p in blanks[keep:]:
        p._element.getparent().remove(p._element)


def balance_sig_table(d):
    t = d.tables[-1]
    col0 = t.rows[1].cells[0]
    # trim the plain 6-blank-line PIHAK PERTAMA sig cell to 4, matching
    # PIHAK KEDUA's {%ttd}/{%stempel} cell (2 tag lines + 4 blank)
    paras = col0.paragraphs
    if len(paras) > 4:
        for p in paras[4:]:
            p._element.getparent().remove(p._element)


def nudge_logo(d):
    hdr = d.sections[0].header
    xml = hdr._element.xml
    if 'wp:posOffset>5057700' in xml:
        return  # idempotent
    ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    for anchor in hdr._element.iter(qn('wp:anchor')):
        posH = anchor.find(qn('wp:positionH'))
        if posH is None:
            continue
        off = posH.find(qn('wp:posOffset'))
        if off is not None and off.text == '4457700':
            off.text = '5057700'  # push ~0.66in further right, clear of centered title/nomor text


def run():
    d = docx.Document(PATH)
    fix_perangkat_jumlah(d)
    trim_blank_paragraphs(d)
    balance_sig_table(d)
    nudge_logo(d)
    d.save(PATH)
    print('revised BAST')


run()

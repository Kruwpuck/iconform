#!/usr/bin/env python3
"""NODIN: drop the "Nomor : {nomor}" row (field removed from the form —
Nota Dinas doesn't need a surat number). Run once from repo root:
python3 scripts/revisi_nodin.py
"""
import docx

d = docx.Document('templates/docx/NODIN.docx')
t = d.tables[0]
if '{nomor}' in t.rows[0].cells[-1].text:
    tbl = t._tbl
    tbl.remove(t.rows[0]._tr)
    print('removed Nomor row')
else:
    print('Nomor row already gone (idempotent no-op)')

d.save('templates/docx/NODIN.docx')

# verify
d2 = docx.Document('templates/docx/NODIN.docx')
for ri, row in enumerate(d2.tables[0].rows):
    print('row', ri, [c.text for c in row.cells])

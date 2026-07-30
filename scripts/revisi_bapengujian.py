#!/usr/bin/env python3
"""Revisi: BA Pengujian — trim excess blank paragraphs so the doc fits one
page (was overflowing to 2 pages purely from empty spacer paragraphs).
Run once from repo root:  python3 scripts/revisi_bapengujian.py
"""
import docx

PATH = 'templates/docx/BA_PENGUJIAN.docx'


def drop_trailing_blanks(container_paras, keep):
    """Remove empty paragraphs from the end of a list, keeping `keep` of them."""
    blanks = [p for p in container_paras if not p.text.strip()]
    to_remove = blanks[keep:] if keep < len(blanks) else []
    for p in to_remove:
        p._element.getparent().remove(p._element)


d = docx.Document(PATH)
body = d.element.body

# 1. Trailing blank paragraphs after the signature table (was 8 — pure page-2 filler)
body_children = list(body)
tbl_indices = [i for i, el in enumerate(body_children) if el.tag.endswith('}tbl')]
last_tbl_idx = tbl_indices[-1]
trailing_blank_ps = [
    el for el in body_children[last_tbl_idx + 1:]
    if el.tag.endswith('}p') and not ''.join(el.itertext()).strip()
]
for el in trailing_blank_ps[1:]:  # keep exactly 1 so sectPr has a paragraph to attach to
    body.remove(el)

# 2. Blank paragraphs between "Demikian pernyataan..." and the signature table (was 3)
first_tbl_idx = tbl_indices[0]
between_blank_ps = [
    el for el in body_children[first_tbl_idx + 1:last_tbl_idx]
    if el.tag.endswith('}p') and not ''.join(el.itertext()).strip()
]
for el in between_blank_ps[1:]:  # keep 1
    body.remove(el)

# 3. Signature-space blank paragraphs inside the sig table cells — trim, don't remove
sig_table = d.tables[-1]
targets = [
    (sig_table.rows[1].cells[0], 3),  # PIHAK PERTAMA sig space (was 7 blanks)
    (sig_table.rows[1].cells[2], 2),  # PIHAK KEDUA sig space, after {%ttd}/{%stempel} (was 4)
    (sig_table.rows[2].cells[0], 1),  # under {namaPihakPertama} (was 2)
    (sig_table.rows[2].cells[2], 2),  # under {namaPihakKedua} (was 4)
]
for cell, keep in targets:
    drop_trailing_blanks(cell.paragraphs, keep)

d.save(PATH)
print('BA_PENGUJIAN.docx trimmed.')

# ── verify ───────────────────────────────────────────────────────────────
d2 = docx.Document(PATH)
body2 = list(d2.element.body)
print('body children:', [el.tag.split("}")[-1] for el in body2])
t2 = d2.tables[-1]
for ri, row in enumerate(t2.rows):
    for ci, cell in enumerate(row.cells):
        print('sig row', ri, 'cell', ci, [p.text for p in cell.paragraphs])

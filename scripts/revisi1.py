#!/usr/bin/env python3
"""Revisi 1: add name tags under sigs, fix BAKL year, logo placeholder,
BAI project-leader section, SURAT_TUGAS YBD block, NODIN multi-row loop.
Run once from repo root:  python3 scripts/revisi1.py
"""
import re
import docx

DOTS = re.compile(r'…[….]*|\.{3,}')


def para_sub(p, pairs):
    if not p.runs:
        return
    text = ''.join(r.text for r in p.runs)
    new = text
    for pat, repl in pairs:
        if isinstance(pat, str):
            new = new.replace(pat, repl)
        else:
            new = pat.sub(repl, new)
    if new != text:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ''


def add_name_para(cell, tag, jabatan_tag=None):
    """Append bold+underline name paragraph (and optional jabatan) to sig cell."""
    if tag in cell.text:
        return  # idempotent
    p = cell.add_paragraph()
    run = p.add_run(f'( {tag} )')
    run.bold = True
    run.underline = True
    if jabatan_tag:
        p2 = cell.add_paragraph()
        p2.add_run(jabatan_tag)


def find_ttd_row_col(table):
    """Return (ri, ci) of the cell containing {%ttd}."""
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if '{%ttd}' in cell.text:
                return ri, ci
    return None, None


def run(name, fn):
    d = docx.Document(f'templates/docx/{name}.docx')
    fn(d)
    d.save(f'templates/docx/{name}.docx')
    print('revised', name)


# ── BA_PENGUJIAN ────────────────────────────────────────────────────────────
def ba_pengujian(d):
    # Logo in header: prepend {%logoMitra} paragraph before title
    hdr = d.sections[0].header
    if '{%logoMitra}' not in hdr.paragraphs[0].text:
        p = hdr.paragraphs[0].insert_paragraph_before('')
        p.add_run('{%logoMitra}')

    # Name tags under sig
    sig = d.tables[-1]
    ri, ci = find_ttd_row_col(sig)
    if ri is not None:
        add_name_para(sig.rows[ri].cells[ci], '{namaPihakKedua}')      # PLN
        add_name_para(sig.rows[ri].cells[0], '{namaPihakPertama}')      # partner


# ── BAP ─────────────────────────────────────────────────────────────────────
def bap(d):
    sig = d.tables[-1]
    ri, ci = find_ttd_row_col(sig)
    if ri is not None:
        add_name_para(sig.rows[ri].cells[ci], '{namaWakilIcon}')        # PLN
        add_name_para(sig.rows[ri].cells[0], '{namaWakilPelanggan}')    # partner


# ── BAST ─────────────────────────────────────────────────────────────────────
def bast(d):
    # Bold the date phrase in intro para
    intro = d.paragraphs[1]
    for run in intro.runs:
        t = run.text
        if any(x in t for x in ('{hari}', '{tanggal}', '{bulan}', '{tahun}',
                                  'hari ini', 'tanggal', 'bulan', 'tahun')):
            run.bold = True

    # Name tags
    sig = d.tables[-1]
    ri, ci = find_ttd_row_col(sig)
    if ri is not None:
        add_name_para(sig.rows[ri].cells[ci], '{namaPihakKedua}',
                      jabatan_tag='{jabatanPihakKedua}')                 # PLN
        add_name_para(sig.rows[ri].cells[0], '{namaPihakPertama}')       # partner


# ── BAI ──────────────────────────────────────────────────────────────────────
def bai(d):
    # Insert "Dan Penyedia Layanan" block before "Demikian berita acara"
    for i, p in enumerate(d.paragraphs):
        if 'Demikian berita acara' in p.text:
            if 'namaProjectLeader' not in d.paragraphs[i - 1].text:
                p.insert_paragraph_before('')
                p.insert_paragraph_before('Nama Project Team Leader\t: {namaProjectLeader}')
                p.insert_paragraph_before('Dan Penyedia Layanan PT PLN ICON PLUS')
            break

    # Name tags under sig
    sig = d.tables[-1]
    ri, ci = find_ttd_row_col(sig)
    if ri is not None:
        add_name_para(sig.rows[ri].cells[ci], '{namaProjectLeader}')    # PLN
        add_name_para(sig.rows[ri].cells[0], '{namaWakil}')             # partner


# ── SURAT_TUGAS ──────────────────────────────────────────────────────────────
def surat_tugas(d):
    # Insert "Yang bertanda tangan" block before "Menugaskan kepada:"
    for p in d.paragraphs:
        if 'Menugaskan kepada:' in p.text:
            if 'namaPemberi' not in p.text:
                p.insert_paragraph_before('')
                p.insert_paragraph_before('Jabatan\t: {jabatanPemberi}')
                p.insert_paragraph_before('Nama\t: {namaPemberi}')
                p.insert_paragraph_before('Yang bertanda tangan dibawah ini:')
            break

    # Name tag after {%stempel}
    for i, p in enumerate(d.paragraphs):
        if '{%stempel}' in p.text:
            if i + 1 >= len(d.paragraphs) or '{namaPemberi}' not in d.paragraphs[i + 1].text:
                # insert after stempel
                stempel_p = p
                new_p = stempel_p.insert_paragraph_before('')  # inserts BEFORE; we want after
                # use XML to insert after
                from docx.oxml import OxmlElement
                np = OxmlElement('w:p')
                nr = OxmlElement('w:r')
                nrpr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                u = OxmlElement('w:u')
                u.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
                nrpr.append(b)
                nrpr.append(u)
                nr.append(nrpr)
                nt = OxmlElement('w:t')
                nt.text = '( {namaPemberi} )'
                nr.append(nt)
                np.append(nr)
                stempel_p._element.addnext(np)
                # remove the blank we accidentally inserted before
                new_p._element.getparent().remove(new_p._element)
            break


# ── BAKL ─────────────────────────────────────────────────────────────────────
def bakl(d):
    # Fix "2026telah" → "{tahun} telah"
    for p in d.paragraphs:
        if '2026' in p.text or 'Tahun ' in p.text:
            para_sub(p, [(re.compile(r'Tahun\s+\d{4}(\w)'), r'Tahun {tahun} \1'),
                         (re.compile(r'2026(\w)'), r'{tahun} \1')])

    # Name tags under sig
    sig = d.tables[-1]
    ri, ci = find_ttd_row_col(sig)
    if ri is not None:
        add_name_para(sig.rows[ri].cells[ci], '{wakilPihakPertama}')    # PLN
        add_name_para(sig.rows[ri].cells[0], '{wakilPihakKedua}')       # partner


# ── NODIN ─────────────────────────────────────────────────────────────────────
def nodin(d):
    # Remove {%ttd} and {%stempel} paragraphs
    to_remove = [p for p in d.paragraphs if p.text.strip() in ('{%ttd}', '{%stempel}')]
    for p in to_remove:
        p._element.getparent().remove(p._element)

    # Tag MI lines (replace full cell text with tag)
    mi_table = d.tables[1]
    mi_table.rows[0].cells[0].paragraphs[0].runs[0].text = '{miDana}'
    for r in mi_table.rows[0].cells[0].paragraphs[0].runs[1:]:
        r.text = ''
    mi_table.rows[1].cells[0].paragraphs[0].runs[0].text = '{miList}'
    for r in mi_table.rows[1].cells[0].paragraphs[0].runs[1:]:
        r.text = ''

    # Convert item row to docxtemplater loop {#items}...{/items}
    items = d.tables[2]
    row2 = items.rows[2]
    # cell[0]: "1" → "{#items}" (loop open)
    para_sub(row2.cells[0].paragraphs[0], [('1', '{#items}')])
    if not row2.cells[0].paragraphs[0].text.strip():
        row2.cells[0].paragraphs[0].add_run('{#items}')
    # cell[5]: "Rp{jumlahTotal}" → "Rp{jumlahTotal}{/items}" (loop close)
    para_sub(row2.cells[5].paragraphs[0],
             [('Rp{jumlahTotal}', 'Rp{jumlahTotal}{/items}')])


run('BA_PENGUJIAN', ba_pengujian)
run('BAP', bap)
run('BAST', bast)
run('BAI', bai)
run('SURAT_TUGAS', surat_tugas)
run('BAKL', bakl)
run('NODIN', nodin)
print('Done.')

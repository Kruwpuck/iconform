"""Stop the BAST signature block from orphaning onto a second page.

The signature table's cells carried 4-5 trailing empty paragraphs each (~1.4in
of invisible padding). They filled the rest of page 1 after the ttd/stempel, so
the names row broke to page 2 — a page holding nothing but the header and two
names.

Fix: drop the trailing empties, then keepNext the rows above the names so the
block can never split again even if the parties' details grow.

Idempotent: re-running finds nothing to strip and the flags already set.
"""
from pathlib import Path

from docx import Document

PATH = Path('templates/docx/BAST.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def main() -> None:
    doc = Document(str(PATH))
    sig = doc.tables[1]
    dropped = 0
    kept = 0

    for row in sig.rows:
        for cell in row.cells:
            # a cell must keep at least one paragraph
            while len(cell.paragraphs) > 1 and not cell.paragraphs[-1].text.strip():
                p = cell.paragraphs[-1]._p
                p.getparent().remove(p)
                dropped += 1

    # every row but the last sticks to the row below it
    for row in sig.rows[:-1]:
        for cell in row.cells:
            for para in cell.paragraphs:
                pPr = para._p.find(W + 'pPr')
                if pPr is None:
                    pPr = para._p.makeelement(W + 'pPr', {})
                    para._p.insert(0, pPr)
                keep = pPr.find(W + 'keepNext')
                if keep is None:
                    keep = pPr.makeelement(W + 'keepNext', {})
                    pPr.insert(0, keep)
                if keep.get(W + 'val') != '1':
                    keep.set(W + 'val', '1')
                    kept += 1

    if not dropped and not kept:
        print('nothing to do — already tidy')
        return

    doc.save(str(PATH))
    print('dropped %d padding paragraphs, keepNext on %d paragraphs' % (dropped, kept))
    check = Document(str(PATH)).tables[1]
    for ri, row in enumerate(check.rows):
        print(' row', ri, [[p.text for p in c.paragraphs] for c in row.cells])


if __name__ == '__main__':
    main()

"""Tidy the signature blocks of BAI and BAKL.

BAI: the left signature cell held an extra empty paragraph the right one did
not (revisi3.py blanked a duplicate instead of removing it), so
"( namaWakil )" printed one line lower than "( namaProjectLeader )".

BAKL: the signature row forced trHeight=1834.7 twips (~1.27in) of blank space,
which both dropped the names far below the TTD and pushed the last block onto a
second page. Trimmed to 900 twips (45pt) — still taller than one mark
(MARK_HEIGHT_PT = 33.75) — plus three redundant blank body paragraphs.

Idempotent: each step checks the current state before touching anything.
"""
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

SIG_ROW_TWIPS = '900'


def drop(el) -> None:
    el.getparent().remove(el)


def fix_bai() -> None:
    path = Path('templates/docx/BAI.docx')
    doc = Document(str(path))
    row = doc.tables[0].rows[2]
    left, right = row.cells[0], row.cells[1]
    changed = False

    # left cell: drop the leading empty paragraph that offsets the name
    while len(left.paragraphs) > 1 and not left.paragraphs[0].text.strip():
        drop(left.paragraphs[0]._p)
        changed = True

    # right cell: leading spaces inside an already-centered paragraph
    first = right.paragraphs[0]
    if first.runs and first.runs[0].text != first.runs[0].text.lstrip():
        first.runs[0].text = first.runs[0].text.lstrip()
        changed = True

    if not changed:
        print('BAI: nothing to do')
        return
    doc.save(str(path))
    check = Document(str(path)).tables[0].rows[2]
    print('BAI:', [[p.text for p in c.paragraphs] for c in check.cells])


def fix_bakl() -> None:
    path = Path('templates/docx/BAKL.docx')
    doc = Document(str(path))
    changed = False

    height = doc.tables[2].rows[1]._tr.find(W + 'trPr').find(W + 'trHeight')
    if height.get(W + 'val') != SIG_ROW_TWIPS:
        print('BAKL: signature row', height.get(W + 'val'), '->', SIG_ROW_TWIPS)
        height.set(W + 'val', SIG_ROW_TWIPS)
        changed = True

    # blank body paragraphs that only pad: before the petugas table, the second
    # of the pair before the signature table, and the trailing one.
    # Highest index first so the earlier ones keep their positions.
    body = list(doc.element.body.iterchildren())
    for i in (25, 23, 17):
        if i >= len(body):
            continue  # already removed on an earlier run
        el = body[i]
        if el.tag == W + 'p' and not Paragraph(el, doc).text.strip():
            drop(el)
            changed = True
        else:
            print('BAKL: body[%d] is not a blank paragraph — left alone' % i)

    if not changed:
        print('BAKL: nothing to do')
        return
    doc.save(str(path))
    print('BAKL: body children now', len(list(Document(str(path)).element.body.iterchildren())))


def slim_bakl_blanks() -> None:
    """Shrink BAKL's empty separator paragraphs from 10pt to 6pt.

    They carry no text, only their paragraph mark's font size, so the spacing
    between sections survives at ~half the height — the rest of the budget
    needed to keep the whole berita acara on one page.
    """
    path = Path('templates/docx/BAKL.docx')
    doc = Document(str(path))
    changed = 0
    for el in doc.element.body.iterchildren():
        if el.tag != W + 'p' or Paragraph(el, doc).text.strip():
            continue
        pPr = el.find(W + 'pPr')
        if pPr is None:
            pPr = el.makeelement(W + 'pPr', {})
            el.insert(0, pPr)
        rPr = pPr.find(W + 'rPr')
        if rPr is None:
            rPr = pPr.makeelement(W + 'rPr', {})
            pPr.append(rPr)
        for name in ('sz', 'szCs'):
            sz = rPr.find(W + name)
            if sz is None:
                sz = rPr.makeelement(W + name, {})
                rPr.append(sz)
            if sz.get(W + 'val') != '12':  # half-points
                sz.set(W + 'val', '12')
                changed += 1

    if not changed:
        print('BAKL blanks: nothing to do')
        return
    doc.save(str(path))
    print('BAKL blanks: shrank', changed, 'size attributes to 6pt')


if __name__ == '__main__':
    fix_bai()
    fix_bakl()
    slim_bakl_blanks()

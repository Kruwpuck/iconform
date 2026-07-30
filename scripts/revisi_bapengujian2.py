#!/usr/bin/env python3
"""Revisi 2 (fork): BA Pengujian — trim to 1 page, align mitra/PLN logos, fix filename prefix."""
import re
import docx

d = docx.Document('templates/docx/BA_PENGUJIAN.docx')

# 1. Trim 3 stray empty body paragraphs (between intro table and closing text) down to 1
paras = d.paragraphs
empties = [p for p in paras[2:5] if p.text.strip() == '']
for p in empties[1:]:
    p._element.getparent().remove(p._element)

# 2. Equalize signature-cell trailing blank lines (row1/row2 col2 had 1-2 extra vs col0)
t = d.tables[1]
row1c2 = t.rows[1].cells[2].paragraphs
# '{%ttd}\n{%stempel}\n\n' -> drop the extra trailing blank paragraph
if len(row1c2) >= 4 and row1c2[-1].text.strip() == '' and row1c2[-2].text.strip() == '':
    row1c2[-1]._element.getparent().remove(row1c2[-1]._element)

row2c2 = t.rows[2].cells[2].paragraphs
if len(row2c2) >= 3 and row2c2[-1].text.strip() == '' and row2c2[-2].text.strip() != '':
    row2c2[-1]._element.getparent().remove(row2c2[-1]._element)

# 3. Header: tighten spacing between logoMitra line and title line, and re-anchor
#    PLN logo's vertical offset to sit level with the (now closer) title line
#    instead of a large negative offset tuned for a taller gap.
hdr = d.sections[0].header
hp = hdr.paragraphs
for p in hp[:2]:
    pf = p.paragraph_format
    pf.space_before = 0
    pf.space_after = 0

xml = hdr._element.xml
m = re.search(r'<wp:posOffset>-152399</wp:posOffset>', hdr._element.xml)
if m:
    for drawing in hdr._element.iter('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV'):
        off = drawing.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset')
        if off is not None and off.text == '-152399':
            off.text = '0'

d.save('templates/docx/BA_PENGUJIAN.docx')
print('BA_PENGUJIAN.docx revised')

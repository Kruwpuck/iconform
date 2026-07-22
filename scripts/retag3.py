"""
retag3.py — one-time DOCX fixes for revisi2:
1. NODIN: add {nomor} before Perihal line
2. BAKL: {kendala1/2/3} flat → {#kendala}{kendala}{/kendala} loop;
         add data row to Petugas Lapangan table with {#petugas}…{/petugas}
3. SURAT_TUGAS: {nama1/2/3} flat → {#petugas}{nama}{/petugas} loop

Idempotent: each step checks target already done before mutating.
"""

import re
import sys
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

TPL = Path(__file__).parent.parent / "templates" / "docx"

# ─── helpers ──────────────────────────────────────────────────────────────────

def para_text(p) -> str:
    return "".join(r.text or "" for r in p.runs)

def set_para_text(p, text: str):
    """Replace all runs with a single run containing text, preserve first run format."""
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""

def clone_para_fmt(src_para, doc):
    """Clone a paragraph element (deep copy for loop template rows)."""
    return copy.deepcopy(src_para._element)

def insert_para_after(ref_para, new_elem):
    ref_para._element.addnext(new_elem)

def make_loop_para(doc, loop_tag: str, existing_para=None):
    """Create a paragraph element containing just loop_tag text, cloning format from existing_para."""
    if existing_para:
        el = clone_para_fmt(existing_para, doc)
        # Clear all runs, set first to loop_tag
        for r in el.findall('.//' + qn('w:r')):
            el.remove(r)
    else:
        el = OxmlElement('w:p')
    r_el = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = loop_tag
    r_el.append(t_el)
    el.append(r_el)
    return el

# ─── 1. NODIN: add {nomor} ────────────────────────────────────────────────────

def fix_nodin(path: Path):
    doc = Document(path)
    # Check idempotent (in tables too)
    all_text = "\n".join(c.text for tbl in doc.tables for row in tbl.rows for c in row.cells)
    all_text += "\n".join(para_text(p) for p in doc.paragraphs)
    if '{nomor}' in all_text:
        print("NODIN: {nomor} already present, skipping")
        return

    # Perihal is in table t0r0. Add a new row at the top with Nomor : {nomor}
    t0 = doc.tables[0]
    tbl_el = t0._tbl
    # Clone header row structure from t0r0
    header_tr = t0.rows[0]._tr
    new_tr = copy.deepcopy(header_tr)
    # Set cells: "Nomor", ": ", "{nomor}"
    tcs = new_tr.findall(qn('w:tc'))
    labels = ['Nomor', ': ', '{nomor}']
    for i, tc in enumerate(tcs):
        for p_el in tc.findall(qn('w:p')):
            tc.remove(p_el)
        p_el = OxmlElement('w:p')
        r_el = OxmlElement('w:r')
        t_el = OxmlElement('w:t')
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = labels[i] if i < len(labels) else ''
        r_el.append(t_el)
        p_el.append(r_el)
        tc.append(p_el)
    # Insert before first row
    tbl_el.insert(list(tbl_el).index(header_tr), new_tr)
    print("NODIN: inserted Nomor : {nomor} row at top of Perihal table")

    doc.save(path)

# ─── 2. BAKL: {kendala1/2/3} → loop; add petugas table row ───────────────────

def fix_bakl(path: Path):
    doc = Document(path)

    # 2a. kendala loop
    done_kendala = any('{#kendala}' in para_text(p) for p in doc.paragraphs)
    if done_kendala:
        print("BAKL: kendala loop already present")
    else:
        # Find kendala1 paragraph
        kendala_paras = [p for p in doc.paragraphs if re.search(r'\{kendala[123]\}', para_text(p))]
        if not kendala_paras:
            print("BAKL: WARNING — kendala paragraphs not found")
        else:
            # Replace first kendala para with loop start + data row + loop end
            first = kendala_paras[0]
            # Insert {#kendala} before first
            open_el = make_loop_para(doc, '{#kendala}', first)
            first._element.addprevious(open_el)
            # Replace first kendala para text with {kendala}
            set_para_text(first, '{kendala}')
            # Remove remaining flat kendala paras
            for p in kendala_paras[1:]:
                p._element.getparent().remove(p._element)
            # Insert {/kendala} after first (now the data row)
            close_el = make_loop_para(doc, '{/kendala}', first)
            first._element.addnext(close_el)
            print("BAKL: kendala1/2/3 → {#kendala}{kendala}{/kendala} loop")

    # 2b. petugas table row
    # Find table with "Perusahaan" header
    target_table = None
    for tbl in doc.tables:
        for row in tbl.rows:
            cells_text = [c.text.strip() for c in row.cells]
            if any('Perusahaan' in ct for ct in cells_text):
                target_table = tbl
                break
        if target_table:
            break

    if target_table is None:
        print("BAKL: WARNING — Petugas Lapangan table not found")
    else:
        # Check if loop already there
        all_text = "\n".join(c.text for row in target_table.rows for c in row.cells)
        if '{#petugas}' in all_text:
            print("BAKL: petugas loop already present in table")
        else:
            # Check if there's already a data row (non-header)
            header_row = target_table.rows[0]
            # Add a new row by cloning the header row structure
            from docx.oxml import OxmlElement as OE
            tbl_el = target_table._tbl

            # Create loop start row
            def make_loop_row_el(text):
                tr = OE('w:tr')
                tc = OE('w:tc')
                # Merge all cols with gridSpan
                tcPr = OE('w:tcPr')
                gridSpan = OE('w:gridSpan')
                gridSpan.set(qn('w:val'), str(len(header_row.cells)))
                tcPr.append(gridSpan)
                tc.append(tcPr)
                p = OE('w:p')
                r = OE('w:r')
                t = OE('w:t')
                t.text = text
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
                return tr

            # Create data row: {#petugas} | fields | {/petugas}
            # We add 3 rows: loop-open, data, loop-close
            open_tr = make_loop_row_el('{#petugas}')
            close_tr = make_loop_row_el('{/petugas}')

            # Data row cloned from header row, fill with field tags
            data_tr = copy.deepcopy(header_row._tr)
            field_tags = ['{perusahaan}', '{nama}', '{noTelp}', '{lokasiKerja}', '']
            for i, cell_el in enumerate(data_tr.findall(qn('w:tc'))):
                # Clear cell text, set tag
                for p_el in cell_el.findall(qn('w:p')):
                    cell_el.remove(p_el)
                p_el = OE('w:p')
                r_el = OE('w:r')
                t_el = OE('w:t')
                t_el.text = field_tags[i] if i < len(field_tags) else ''
                r_el.append(t_el)
                p_el.append(r_el)
                cell_el.append(p_el)

            tbl_el.append(open_tr)
            tbl_el.append(data_tr)
            tbl_el.append(close_tr)
            print("BAKL: petugas loop rows added to table")

    doc.save(path)

# ─── 3. SURAT_TUGAS: {nama1/2/3} → {#petugas}{nama}{/petugas} ───────────────

def fix_surat_tugas(path: Path):
    doc = Document(path)

    done = any('{#petugas}' in para_text(p) for p in doc.paragraphs)
    if done:
        print("SURAT_TUGAS: petugas loop already present")
        doc.save(path)
        return

    # Find paragraphs containing {nama1}, {nama2}, {nama3}
    nama_paras = [p for p in doc.paragraphs if re.search(r'\{nama[123]\}', para_text(p))]
    if not nama_paras:
        print("SURAT_TUGAS: WARNING — {nama1/2/3} paragraphs not found")
        return

    # Replace first with loop, remove rest
    first = nama_paras[0]
    # Insert {#petugas} before
    open_el = make_loop_para(doc, '{#petugas}', first)
    first._element.addprevious(open_el)
    # Set data row text (keep "Sdr." prefix)
    set_para_text(first, 'Sdr. {nama}')
    # Remove other flat nama paras
    for p in nama_paras[1:]:
        p._element.getparent().remove(p._element)
    # Insert {/petugas} after
    close_el = make_loop_para(doc, '{/petugas}', first)
    first._element.addnext(close_el)
    print("SURAT_TUGAS: {nama1/2/3} → {#petugas}Sdr. {nama}{/petugas} loop")

    doc.save(path)

# ─── main ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("=== NODIN ===")
    fix_nodin(TPL / "NODIN.docx")
    print("\n=== BAKL ===")
    fix_bakl(TPL / "BAKL.docx")
    print("\n=== SURAT_TUGAS ===")
    fix_surat_tugas(TPL / "SURAT_TUGAS.docx")
    print("\nDone.")

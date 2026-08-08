#!/usr/bin/env python3
"""Revisi 3: SURAT_TUGAS Nama/Jabatan alignment + petugas hanging indent,
BAI/UID_JABAR bold intro + instansiPelanggan + customer ttd2/stempel2 +
SN column alignment, BAKL bold intro + drop stray blank petugas rows +
customer ttd2/stempel2 + signature blank line. Idempotent — safe to re-run.
Run once from repo root:  python3 scripts/revisi3.py
"""
import copy
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx


def run(name, fn):
    d = docx.Document(f'templates/docx/{name}.docx')
    fn(d)
    d.save(f'templates/docx/{name}.docx')
    print('revised', name)


def set_tabs_ind(p, tabs, left):
    """Overwrite the w:tabs and w:ind/@left of a paragraph's pPr."""
    pPr = p._p.get_or_add_pPr()
    for tag in ('tabs', 'ind'):
        el = pPr.find(qn('w:' + tag))
        if el is not None:
            pPr.remove(el)
    tabsEl = OxmlElement('w:tabs')
    for pos in tabs:
        t = OxmlElement('w:tab')
        t.set(qn('w:val'), 'left')
        t.set(qn('w:pos'), str(pos))
        tabsEl.append(t)
    indEl = OxmlElement('w:ind')
    indEl.set(qn('w:left'), str(left))
    indEl.set(qn('w:right'), '0')
    indEl.set(qn('w:firstLine'), '0')
    # insert before rPr if present, else append
    rPr = pPr.find(qn('w:rPr'))
    if rPr is not None:
        rPr.addprevious(tabsEl)
        rPr.addprevious(indEl)
    else:
        pPr.append(tabsEl)
        pPr.append(indEl)


def bold_keywords(p, keywords):
    """Split a paragraph's single/mixed runs into segments, bolding any
    segment that contains one of `keywords` (substrings)."""
    text = ''.join(r.text for r in p.runs)
    if not text:
        return
    # already split with correct bold state? idempotency check below
    pattern = '(' + '|'.join(re.escape(k) for k in keywords) + ')'
    parts = re.split(pattern, text)
    if len(parts) <= 1:
        return
    # take formatting template from the first existing run
    template_rPr = p.runs[0]._r.find(qn('w:rPr'))
    # remove all existing runs
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        if template_rPr is not None:
            run._r.get_or_add_rPr()
            # copy font props from template (font, size, color) minus bold
            for child in template_rPr:
                if child.tag == qn('w:b'):
                    continue
                run._r.get_or_add_rPr().append(copy.deepcopy(child))
        run.bold = part in keywords


def blank_dupe_para(cell, before_text_contains):
    """Insert a blank paragraph directly above the paragraph whose text
    contains `before_text_contains`, unless one already exists there."""
    for i, p in enumerate(cell.paragraphs):
        if before_text_contains in p.text:
            if i > 0 and cell.paragraphs[i - 1].text.strip() == '':
                return  # already has a blank line above — idempotent
            p.insert_paragraph_before('')
            return


# ── SURAT_TUGAS ──────────────────────────────────────────────────────────
def surat_tugas(d):
    paras = d.paragraphs
    nama_p = next(p for p in paras if p.text.strip() == 'Nama\t: {namaPemberi}')
    jabatan_pPr = nama_p._p.get_or_add_pPr()
    tabs = [t.get(qn('w:pos')) for t in jabatan_pPr.findall(qn('w:tabs') + '/' + qn('w:tab'))]
    if not tabs:
        tabs_el = jabatan_pPr.find(qn('w:tabs'))
        tabs = [t.get(qn('w:pos')) for t in tabs_el.findall(qn('w:tab'))] if tabs_el is not None else ['1557']
    ind_el = jabatan_pPr.find(qn('w:ind'))
    left = ind_el.get(qn('w:left')) if ind_el is not None else '117'

    loop_open = next(p for p in paras if p.text.strip() == '{#petugas}')
    # 1) "Nama\t:" label line before the petugas loop (idempotent)
    prev_idx = paras.index(loop_open) - 1
    if paras[prev_idx].text.strip() != 'Nama\t:':
        new_p = loop_open.insert_paragraph_before('Nama\t:')
        set_tabs_ind(new_p, [int(t) for t in tabs], int(left))

    # 2) hanging indent for the "Sdr. {nama}" loop row so multi-row lists
    #    align under the value column instead of the left margin
    sdr_p = next(p for p in paras if p.text.strip() == '- Sdr. {nama}')
    set_tabs_ind(sdr_p, [], int(left) + 1670)  # ≈ right after "Nama\t: " column

    # 3) normalize "Tanggal Tugas" / "Lokasi" tab stops + strip the extra
    #    literal spaces that were faking alignment
    for p in paras:
        t = p.text
        if t.startswith('Tanggal Tugas\t'):
            p.runs[0].text = 'Tanggal Tugas\t: {tanggalTugas} s/d {tanggalTugasSelesai}'
            for r in p.runs[1:]:
                r.text = ''
            set_tabs_ind(p, [int(t) for t in tabs], int(left))
        elif t.startswith('Lokasi\t'):
            p.runs[0].text = 'Lokasi\t: {lokasi}'
            for r in p.runs[1:]:
                r.text = ''
            set_tabs_ind(p, [int(t) for t in tabs], int(left))

    # 4) footer color bar — NOT attempted. This template has no header/footer
    #    XML parts at all (word/header*.xml / word/footer*.xml are absent);
    #    its single top graphic is a floating image anchored inside the body
    #    (near "Surat Tugas ini dibuat..."), and no template in this repo
    #    has a distinct bottom/footer bar asset to copy from. Fabricating a
    #    new header/footer part + graphic blind (no local PDF render to
    #    verify) risks corrupting the docx, so this is flagged for manual
    #    fix in Word rather than attempted here.


# ── BAI / UID_JABAR (identical treatment) ──────────────────────────────────
def bai_like(d):
    intro = d.paragraphs[1]
    bold_keywords(intro, ['{hari}', '{tanggal}', '{bulan}', '{tahun}', 'BAIK'])

    sig = d.tables[0]
    # bold {instansiPelanggan} header cell
    cell0 = sig.rows[0].cells[0]
    for r in cell0.paragraphs[0].runs:
        if '{instansiPelanggan}' in r.text:
            r.bold = True

    # customer ttd2/stempel2 (mirrors {%ttd}/{%stempel} already in cell[1])
    ttd_cell = sig.rows[1].cells[0]
    if '{%ttd2}' not in ttd_cell.text:
        p1 = ttd_cell.paragraphs[0] if ttd_cell.paragraphs[0].text == '' else ttd_cell.add_paragraph()
        p1.add_run('{%ttd2}')
        ttd_cell.add_paragraph().add_run('{%stempel2}')

    # blank signature-space line above "( {namaWakil} )"
    blank_dupe_para(sig.rows[2].cells[0], '{namaWakil}')

    # SN column alignment: give "Nama Perangkat" / "Nama Perangkat POP"
    # rows a second, shared tab stop so "SN:" lines up between them.
    for p in d.paragraphs:
        if p.text.startswith('Nama Perangkat\t:') or p.text.startswith('Nama Perangkat POP\t:'):
            pPr = p._p.get_or_add_pPr()
            tabsEl = pPr.find(qn('w:tabs'))
            if tabsEl is not None and len(tabsEl.findall(qn('w:tab'))) < 2:
                t = OxmlElement('w:tab')
                t.set(qn('w:val'), 'left')
                t.set(qn('w:pos'), '6096')
                tabsEl.append(t)

    # widen the shared first tab stop slightly so the longest label
    # ("Nomor Telp/HP, Email") still lands its colon at the same column
    # as the shorter rows instead of overflowing to a default tab stop
    for p in d.paragraphs:
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            continue
        tabsEl = pPr.find(qn('w:tabs'))
        if tabsEl is None:
            continue
        first_tab = tabsEl.find(qn('w:tab'))
        if first_tab is not None and first_tab.get(qn('w:pos')) == '2880':
            first_tab.set(qn('w:pos'), '3020')


# ── BAKL ─────────────────────────────────────────────────────────────────────
def bakl(d):
    intro = d.paragraphs[3]
    bold_keywords(intro, ['{hari}', '{tanggal}', '{bulan}', '{tahun}'])

    # drop the two always-blank filler rows in the Petugas Lapangan table
    # (rows sit between the header row and the {#petugas} loop row; they
    # showed up as stray blank rows on every render, and grew the gap
    # further whenever the user added more loop rows in the UI)
    petugas_tbl = d.tables[1]
    for row in list(petugas_tbl.rows[1:3]):
        if all(c.text.strip() == '' for c in row.cells):
            row._element.getparent().remove(row._element)

    sig = d.tables[2]
    # customer ttd2/stempel2 (mirrors {%ttd}/{%stempel} already in cell[2])
    ttd_cell = sig.rows[1].cells[0]
    if '{%ttd2}' not in ttd_cell.text:
        p1 = ttd_cell.paragraphs[0] if ttd_cell.paragraphs[0].text == '' else ttd_cell.add_paragraph()
        p1.add_run('{%ttd2}')
        ttd_cell.add_paragraph().add_run('{%stempel2}')

    # blank signature-space line above both parenthesized names
    blank_dupe_para(sig.rows[2].cells[0], '{wakilPihakKedua}')
    blank_dupe_para(sig.rows[2].cells[2], '{wakilPihakPertama}')

    # logo placement — investigated: the PLN Icon Plus logo lives in a true
    # header part (word/header*.xml), referenced once via sectPr, with no
    # titlePg override — i.e. it already repeats identically on every page
    # and is not body-flow content. It is not what's shifting. The visible
    # "logo pushed down" effect was the stray blank table rows above
    # inflating the page-1 content height; removed those. No header/logo
    # XML change made (none was warranted).


run('SURAT_TUGAS', surat_tugas)
run('BAI', bai_like)
run('UID_JABAR', bai_like)
run('BAKL', bakl)
print('Done.')

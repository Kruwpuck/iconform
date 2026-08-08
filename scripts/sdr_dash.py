"""Prefix "- " to the "Sdr. {nama}" line in Surat Tugas.

Idempotent: skips paragraphs already starting with "- ". Follows the
load/mutate/save/verify pattern of the other revisi scripts.
"""
from pathlib import Path

from docx import Document

PATH = Path('templates/docx/SURAT_TUGAS.docx')


def main() -> None:
    doc = Document(str(PATH))
    changed = 0
    for p in doc.paragraphs:
        if 'Sdr' not in p.text or p.text.lstrip().startswith('- '):
            continue
        run = p.runs[0]
        run.text = '- ' + run.text
        changed += 1

    if not changed:
        print('nothing to do — already prefixed')
        return

    doc.save(str(PATH))

    # verify
    check = Document(str(PATH))
    for p in check.paragraphs:
        if 'Sdr' in p.text:
            print('now:', repr(p.text))
    print(f'paragraphs changed: {changed}')


if __name__ == '__main__':
    main()

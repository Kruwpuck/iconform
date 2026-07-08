#!/usr/bin/env python3
"""One-time: add {%ttd} + {%stempel} image tags at the PLN signature spot
of every master (in place). Rendered by docxtemplater-image-module-free;
empty upload renders as an invisible 1x1 image.

One tag per paragraph — the free image module renders only the first
image tag when two share a paragraph."""
import docx


def in_cell(doc, row, col, new_para=False):
    cell = doc.tables[-1].rows[row].cells[col]
    if new_para:  # target cell already holds heading text
        cell.add_paragraph().add_run('{%ttd}')
        cell.add_paragraph().add_run('{%stempel}')
    else:
        ps = cell.paragraphs
        ps[0].add_run('{%ttd}')
        (ps[1] if len(ps) > 1 else cell.add_paragraph()).add_run('{%stempel}')
    strip_empty_cell_paras(doc.tables[-1])


def strip_empty_cell_paras(table):
    """Drop the leftover blank lines that reserved manual-signature space —
    with images they push the stempel (and the row) off the page."""
    for row in table.rows:
        for cell in row.cells:
            ps = cell.paragraphs
            empties = [p for p in ps if not p.text.strip()]
            keep_one = len(empties) == len(ps)  # a cell needs >=1 paragraph
            for p in empties[1 if keep_one else 0:]:
                if p._p.getparent() is not None:  # merged cells revisit
                    p._p.getparent().remove(p._p)


def after_para(doc, needle, offset, append=False):
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            doc.paragraphs[i + offset].add_run('{%ttd}')
            if append:
                doc.add_paragraph().add_run('{%stempel}')
            else:
                doc.paragraphs[i + offset + 1].add_run('{%stempel}')
            return
    raise SystemExit(f'needle not found: {needle}')


def run(name, fn):
    d = docx.Document(f'templates/docx/{name}.docx')
    fn(d)
    d.save(f'templates/docx/{name}.docx')
    print('tagged', name)


run('SURAT_TUGAS', lambda d: after_para(d, 'Bandung, {tanggalSurat}', 2))
run('BAI', lambda d: in_cell(d, 1, 1))            # under "PT PLN ICON PLUS"
run('BAKL', lambda d: in_cell(d, 1, 2))           # under "PT PLN ICON PLUS"
run('BAP', lambda d: in_cell(d, 0, 1, new_para=True))  # under "PT PLN ICONPLUS"
run('BA_PENGUJIAN', lambda d: in_cell(d, 1, 2))   # under "PIHAK KEDUA" (PLN)
run('BAST', lambda d: in_cell(d, 1, 2))           # under "PIHAK KEDUA" (PLN)
run('NODIN', lambda d: after_para(d, 'Demikian disampaikan', 1, append=True))

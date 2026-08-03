"""Fix BAKL petugas table loop — docxtemplater row loop needs {#petugas} and
{/petugas} in the SAME row as the data, else the marker-only rows render as
empty rows and break the table borders."""

from pathlib import Path
from docx import Document

TPL = Path(__file__).parent.parent / "templates" / "docx"


def fix_bakl_table():
    doc = Document(TPL / "BAKL.docx")
    for tbl in doc.tables:
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if 'Perusahaan' not in hdr or 'Paraf' not in hdr:
            continue

        rows = tbl.rows
        # locate marker rows
        open_idx = next((i for i, r in enumerate(rows) if '{#petugas}' in r.cells[0].text), None)
        close_idx = next((i for i, r in enumerate(rows) if '{/petugas}' in r.cells[0].text), None)
        data_idx = next((i for i, r in enumerate(rows) if '{perusahaan}' in r.cells[0].text), None)
        if None in (open_idx, close_idx, data_idx):
            print("BAKL: markers not found — already fixed?")
            return

        data_row = rows[data_idx]
        first_cell = data_row.cells[0]
        last_cell = data_row.cells[-1]

        # prepend {#petugas} to first cell, append {/petugas} to last cell
        fp = first_cell.paragraphs[0]
        if fp.runs:
            fp.runs[0].text = '{#petugas}' + fp.runs[0].text
        else:
            fp.add_run('{#petugas}')

        lp = last_cell.paragraphs[0]
        if lp.runs:
            lp.runs[-1].text = lp.runs[-1].text + '{/petugas}'
        else:
            lp.add_run('{/petugas}')

        # remove the two marker-only rows (delete higher index first)
        for idx in sorted([open_idx, close_idx], reverse=True):
            tbl._tbl.remove(rows[idx]._tr)

        doc.save(TPL / "BAKL.docx")
        print("BAKL: petugas loop markers merged into data row, empty rows removed")
        return

    print("BAKL: petugas table not found")


if __name__ == '__main__':
    fix_bakl_table()

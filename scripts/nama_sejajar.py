"""Line the Surat Tugas petugas names up under one column.

Before: "Nama<tab>:" on its own line, then a {#petugas} paragraph loop putting
every "- Sdr. {nama}" on a separate, fixed-indent paragraph — so the first name
sat *below* the label instead of next to it.

After: one paragraph "Nama<tab>:<tab>{daftarPetugas}" with a hanging indent, so
the first name sits on the label line and every following name (fed as a "\n"
separated block by fillDocx) wraps to exactly the same column.

Idempotent: bails out if {daftarPetugas} is already there.
"""
from pathlib import Path

from docx import Document

PATH = Path('templates/docx/SURAT_TUGAS.docx')

LABEL = 'Nama'
TAG = '{daftarPetugas}'
NAME_POS = 1787  # twips — the old "- Sdr." paragraph indent; names align here


def main() -> None:
    doc = Document(str(PATH))
    paras = doc.paragraphs

    if any(TAG in p.text for p in paras):
        print('nothing to do — already converted')
        return

    label_i = next(i for i, p in enumerate(paras) if p.text.startswith(LABEL + '\t:') and i > 14)
    loop = [i for i, p in enumerate(paras) if p.text.strip() in ('{#petugas}', '{/petugas}')]
    name_i = next(i for i, p in enumerate(paras) if '- Sdr. {nama}' in p.text)
    assert len(loop) == 2 and loop[0] < name_i < loop[1], 'unexpected petugas loop layout'

    p = paras[label_i]
    ns = p._p.nsmap['w']
    q = lambda n: '{%s}%s' % (ns, n)

    # hanging indent: first line back at the old left margin, every wrapped
    # line at the name column
    pPr = p._p.find(q('pPr'))
    ind = pPr.find(q('ind'))
    left = int(ind.get(q('left')))
    ind.set(q('left'), str(NAME_POS))
    ind.set(q('hanging'), str(NAME_POS - left))
    if ind.get(q('firstLine')) is not None:
        del ind.attrib[q('firstLine')]

    # explicit tab stop at the name column (Word implies one at a hanging
    # indent, LibreOffice is happier being told)
    tabs = pPr.find(q('tabs'))
    tabs.append(tabs.makeelement(q('tab'), {q('val'): 'left', q('pos'): str(NAME_POS)}))

    # "Nama" <tab> ":" -> "Nama" <tab> ":" <tab> "{daftarPetugas}"
    run = p.runs[-1]._r
    run.append(run.makeelement(q('tab'), {}))
    t = run.makeelement(q('t'), {})
    t.text = TAG
    run.append(t)

    # drop the now-dead loop paragraphs
    for i in sorted({loop[0], name_i, loop[1]}, reverse=True):
        el = paras[i]._p
        el.getparent().remove(el)

    doc.save(str(PATH))

    check = Document(str(PATH))
    for cp in check.paragraphs:
        if 'Nama' in cp.text or 'petugas' in cp.text or 'Sdr' in cp.text:
            print(repr(cp.text), '| ind=', cp.paragraph_format.left_indent)


if __name__ == '__main__':
    main()

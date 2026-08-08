"""Make wrapped values line up, and unstick the BA Pengujian signature block.

1. UID_JABAR / BAI print their specs as "Label<tab>: {value}". Only Originating
   and Terminating had a hanging indent, so every other long value (Nama
   Pelanggan, Alamat Kantor, ...) wrapped back to the left margin and collided
   with the label column. Each such paragraph now hangs at its own first tab
   stop, so continuation lines sit exactly under the value.

2. BA_PENGUJIAN's closing paragraph had space_after=0 with the signature table
   right beneath it. Given the same 18pt the paragraph above it already uses.

Idempotent: both steps compare before writing.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Emu

LABEL_VALUE = re.compile(r'^[^\t]{1,30}\t\s*:')
HANGING_FILES = ['UID_JABAR.docx', 'BAI.docx']
SPACE_AFTER = Emu(228600)  # 18pt — matches the preceding paragraph


def hang_label_values(name: str) -> None:
    path = Path('templates/docx') / name
    doc = Document(str(path))
    changed = 0
    for p in doc.paragraphs:
        stops = p.paragraph_format.tab_stops
        if not LABEL_VALUE.match(p.text) or not len(stops):
            continue
        col = stops[0].position  # the value column this line already tabs to
        pf = p.paragraph_format
        if pf.left_indent == col and pf.first_line_indent == -col:
            continue
        pf.left_indent = col
        pf.first_line_indent = -col
        changed += 1

    if not changed:
        print(name, '- nothing to do')
        return
    doc.save(str(path))
    print(name, '- hanging indent on', changed, 'label lines')


def space_before_signature() -> None:
    path = Path('templates/docx/BA_PENGUJIAN.docx')
    doc = Document(str(path))
    target = next(p for p in doc.paragraphs if p.text.startswith('Demikian pernyataan'))
    if target.paragraph_format.space_after == SPACE_AFTER:
        print('BA_PENGUJIAN - nothing to do')
        return
    target.paragraph_format.space_after = SPACE_AFTER
    doc.save(str(path))
    print('BA_PENGUJIAN - space_after set before the signature table')


if __name__ == '__main__':
    for f in HANGING_FILES:
        hang_label_values(f)
    space_before_signature()

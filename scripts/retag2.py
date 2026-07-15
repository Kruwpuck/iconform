#!/usr/bin/env python3
"""One-time cleanup: sisa literal contoh / dotted placeholder yang lolos retag.py.

retag.py punya guard `if '{' in para.text: return` — paragraf yang sudah
bertag di-skip padahal masih menyimpan dots/literal lain. Fix 5 titik di
BAI, UID_JABAR, BAST. Idempotent: skip kalau target sudah tidak ada.
"""
import re
from pathlib import Path
from docx import Document

DOTS = re.compile(r'[……\.]{2,}')
DST = Path('templates/docx')

UID_LAYANAN = ('Managed Service Furniture Kantor Yogyakarta dan Depo KRL Solo Jebres '
               'Tahun 2025-2027 Area VI Yogyakarta Division PT Kereta Commuter Indonesia')
UID_PELANGGAN = 'PT. PLN (PERSERO) UNIT INDUK DISTRIBUSI JAWA BARAT'
UID_GEDUNG = re.compile(r'GEDUNG BALAI SUMUR BANDUNG[^\t\n]*')


def merge(para, new_text):
    if not para.runs:
        para.add_run(new_text)
    else:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''


def sub_dots(para, *tags):
    """Ganti tiap run dots berikutnya dengan tag — tanpa guard tag."""
    text = para.text
    changed = False
    for tag in tags:
        if DOTS.search(text):
            text = DOTS.sub(tag, text, count=1)
            changed = True
    if changed:
        merge(para, text)
    return changed


def do_bai():
    doc = Document(DST / 'BAI.docx')
    ok = sub_dots(doc.paragraphs[1], '{namaLayanan}', '{namaPelanggan}')
    ok |= sub_dots(doc.paragraphs[12], '{terminating}')
    if ok:
        doc.save(DST / 'BAI.docx')
    print(f'BAI: {"updated" if ok else "already clean"}')


def do_uid_jabar():
    doc = Document(DST / 'UID_JABAR.docx')
    ok = False
    p1 = doc.paragraphs[1]
    if UID_LAYANAN in p1.text:
        merge(p1, p1.text.replace(UID_LAYANAN, '{namaLayanan}')
                        .replace(UID_PELANGGAN, '{namaPelanggan}'))
        ok = True
    p12 = doc.paragraphs[12]
    if UID_GEDUNG.search(p12.text):
        merge(p12, UID_GEDUNG.sub('{terminating}', p12.text).rstrip())
        ok = True
    if ok:
        doc.save(DST / 'UID_JABAR.docx')
    print(f'UID_JABAR: {"updated" if ok else "already clean"}')


def do_bast():
    doc = Document(DST / 'BAST.docx')
    ok = False
    cell = doc.tables[0].rows[2].cells[3]
    for para in cell.paragraphs:
        if DOTS.search(para.text) and '{' not in para.text:
            merge(para, '')
            ok = True
    if ok:
        doc.save(DST / 'BAST.docx')
    print(f'BAST: {"updated" if ok else "already clean"}')


if __name__ == '__main__':
    do_bai()
    do_uid_jabar()
    do_bast()
